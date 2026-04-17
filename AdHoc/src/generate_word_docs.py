#!/usr/bin/env python3
"""
Generate Word documents with email sequences — one .docx per contact.
Styled like an email client thread view.
"""

import json
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ADHOC_ROOT = os.path.dirname(_SCRIPT_DIR)
_PROJECT_ROOT = os.path.dirname(_ADHOC_ROOT)

SENDER_NAME = "Tomasz Uściński"
SENDER_TITLE = "Senior Client Partner | Procurement Technology"
SENDER_COMPANY = "PROFITIA Management Consultants"
SENDER_PHONE = "+48 787 417 293"
SENDER_EMAIL_ADDR = "tomasz.uscinski@profitia.pl"

SIGNATURE_BLOCK = (
    f"Pozdrawiam serdecznie,\n\n"
    f"{SENDER_NAME}\n"
    f"{SENDER_TITLE}\n"
    f"{SENDER_PHONE}\n\n"
    f"{SENDER_COMPANY}\n"
    f"SpendGuru Data-driven Procurement | Certyfikowany Partner CIPS w Polsce | Doradztwo | Szkolenia | Analityka zakupowa\n"
    f"02-715 Warszawa, Villa Metro, ul. Puławska 145, V p."
)


def add_email_block(doc, step_label, subject, body, recipient_email, is_reply=False):
    """Add one email block to the document, styled like an email client."""

    # Separator line before replies
    if is_reply:
        sep = doc.add_paragraph()
        sep_run = sep.add_run("─" * 60)
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        sep.space_before = Pt(16)
        sep.space_after = Pt(4)

    # Step label (Email 1 / Follow-up 1 / Follow-up 2)
    label_p = doc.add_paragraph()
    label_run = label_p.add_run(step_label)
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    label_p.space_after = Pt(2)

    # Header block (From / To / Subject)
    header_style = {"size": Pt(9), "color": RGBColor(0x66, 0x66, 0x66)}

    def add_header_line(label_text, value_text):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        lbl = p.add_run(f"{label_text}: ")
        lbl.bold = True
        lbl.font.size = header_style["size"]
        lbl.font.color.rgb = header_style["color"]
        val = p.add_run(value_text)
        val.font.size = header_style["size"]
        val.font.color.rgb = header_style["color"]

    re_prefix = "Re: " if is_reply else ""
    add_header_line("Od", f"{SENDER_NAME} <{SENDER_EMAIL_ADDR}>")
    add_header_line("Do", recipient_email)
    add_header_line("Temat", f"{re_prefix}{subject}")

    # Spacer
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(4)
    spacer.space_after = Pt(4)

    # Body text
    # Check if body already contains signature-like closing
    body_clean = body.strip()
    has_signature = any(s in body_clean.lower() for s in [
        "pozdrawiam serdecznie",
        "tomasz uściński",
    ])

    for para_text in body_clean.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        # Handle single newlines within paragraph
        para_text = para_text.replace("\n", "\n")
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.space_after = Pt(4)

    # Add signature if body doesn't already have one
    if not has_signature:
        sig_spacer = doc.add_paragraph()
        sig_spacer.space_before = Pt(8)
        sig_spacer.space_after = Pt(0)

        for sig_line in SIGNATURE_BLOCK.split("\n"):
            p = doc.add_paragraph()
            p.space_before = Pt(0)
            p.space_after = Pt(0)
            run = p.add_run(sig_line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.name = "Calibri"


def generate_doc_for_contact(contact: dict, output_dir: str):
    """Generate a single Word document with 3-email sequence for a contact."""

    name = contact.get("name", "Unknown")
    company = contact.get("apollo_company") or contact.get("company", "")
    title = contact.get("apollo_title") or contact.get("known_title", "")
    email = contact.get("apollo_email", "")
    panel = contact.get("panel_title", "")

    email_1_subject = contact.get("email_1_subject", "")
    email_1_body = contact.get("email_1_body", "")
    email_2_subject = contact.get("email_2_subject", "")
    email_2_body = contact.get("email_2_body", "")
    email_3_subject = contact.get("email_3_subject", "")
    email_3_body = contact.get("email_3_body", "")

    if not email_1_body:
        return None

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_p = doc.add_heading(level=1)
    title_run = title_p.add_run(f"{name}")
    title_run.font.size = Pt(16)

    # Contact info
    info_p = doc.add_paragraph()
    info_parts = []
    if title:
        info_parts.append(title)
    if company:
        info_parts.append(company)
    if email:
        info_parts.append(email)
    info_run = info_p.add_run(" | ".join(info_parts))
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if panel:
        panel_p = doc.add_paragraph()
        panel_run = panel_p.add_run(f"Panel: {panel}")
        panel_run.font.size = Pt(9)
        panel_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        panel_run.italic = True

    # Thin line
    doc.add_paragraph().space_after = Pt(2)

    # Email 1
    add_email_block(
        doc,
        step_label="EMAIL 1 — Pierwszy kontakt",
        subject=email_1_subject,
        body=email_1_body,
        recipient_email=email,
        is_reply=False,
    )

    # Email 2
    if email_2_body:
        add_email_block(
            doc,
            step_label="EMAIL 2 — Follow-up",
            subject=email_2_subject or email_1_subject,
            body=email_2_body,
            recipient_email=email,
            is_reply=True,
        )

    # Email 3
    if email_3_body:
        add_email_block(
            doc,
            step_label="EMAIL 3 — Final follow-up",
            subject=email_3_subject or email_1_subject,
            body=email_3_body,
            recipient_email=email,
            is_reply=True,
        )

    # Save
    safe_name = name.replace(" ", "_").lower()
    safe_name = "".join(c for c in safe_name if c.isalnum() or c in "_-")
    filename = f"{safe_name}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def main():
    # Find the latest output run
    outputs_dir = os.path.join(_ADHOC_ROOT, "outputs")
    runs = sorted([d for d in os.listdir(outputs_dir) if os.path.isdir(os.path.join(outputs_dir, d))])
    if not runs:
        print("[ERROR] No output runs found.")
        sys.exit(1)

    latest_run = runs[-1]
    run_dir = os.path.join(outputs_dir, latest_run)
    results_path = os.path.join(run_dir, "campaign_results.json")

    if not os.path.exists(results_path):
        print(f"[ERROR] campaign_results.json not found in {run_dir}")
        sys.exit(1)

    # Create DOCS directory
    docs_dir = os.path.join(run_dir, "DOCS")
    os.makedirs(docs_dir, exist_ok=True)

    # Load results
    with open(results_path, "r", encoding="utf-8") as f:
        results = json.load(f)

    ready = [r for r in results if r.get("status") == "ready_for_review"]
    print(f"Generating Word docs for {len(ready)} contacts...")

    generated = []
    for contact in ready:
        path = generate_doc_for_contact(contact, docs_dir)
        if path:
            print(f"  ✓ {os.path.basename(path)}")
            generated.append(path)

    print(f"\nDone. {len(generated)} docs saved to:\n  {docs_dir}")


if __name__ == "__main__":
    main()
