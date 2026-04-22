#!/usr/bin/env python3
"""
Generate email sequences via LLM for 10 contacts without Apollo email addresses,
then save as v3 (Outlook-ready) DOCX files.

Pipeline:
1. Load 10 no-email contacts from campaign_results.json
2. Fix vocative forms (campaign_results.json has base forms, not vocative)
3. Call LLM (centralny llm_router → OPENAI_PRIMARY_MODEL) with adhoc_email_writer.md prompt + context
4. Apply global rules (em dash → hyphen, CTA gender-aware, role phrasing)
5. Generate v3 DOCX files (Outlook-ready format)
"""

import json
import os
import sys
import unicodedata

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ADHOC_ROOT = os.path.dirname(_SCRIPT_DIR)
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(_ADHOC_ROOT)))

# Add paths
sys.path.insert(0, os.path.join(_PROJECT_ROOT, "src"))

from llm_client import generate_json, is_llm_available
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Correct vocative forms ──────────────────────────────────────
VOCATIVE_FIX = {
    "Tomasz": "Tomaszu",
    "Barbara": "Barbaro",
    "Rafał": "Rafale",
    "Marcin": "Marcinie",
    "Piotr": "Piotrze",
    "Wojciech": "Wojciechu",
    "Krzysztof": "Krzysztofie",
    "Agnieszka": "Agnieszko",
}

# ── Signature spec ──────────────────────────────────────────────
SIG_LINES = [
    {"text": "Tomasz Uściński", "size": 11, "bold": True, "color": "000000"},
    {"text": "Senior Client Partner | Procurement Technology", "size": 11, "bold": False, "color": "000000"},
    {"text": "+48 787 417 293", "size": 11, "bold": False, "color": "000000"},
    {"text": "", "size": 11, "bold": False, "color": "000000"},
    {"text": "PROFITIA Management Consultants Mazurowski i Wspólnicy Spółka Jawna", "size": 11, "bold": True, "color": "000000"},
    {"text": "SpendGuru Data-driven Procurement | Certyfikowany Partner CIPS w Polsce | Doradztwo | Szkolenia | Analityka zakupowa", "size": 11, "bold": False, "color": "000000"},
    {"text": "02-715 Warszawa, Villa Metro, ul. Puławska 145, V p.", "size": 11, "bold": False, "color": "000000"},
    {"text": "", "size": 11, "bold": False, "color": "000000"},
    {"text": "Uwaga: Ten e-mail jest poufny i przeznaczony tylko dla adresata (-ów) tej wiadomości. Jeżeli nie jest Pan/Pani jego zamierzonym adresatem, to jakiekolwiek posługiwanie się, kopiowanie lub rozpowszechnianie tego e-maila jest zabronione. Jeśli otrzymał(a) Pan/Pani ten e-mail omyłkowo, proszę o natychmiastowy kontakt z nadawcą i usunięcie tego e-maila.", "size": 9, "bold": False, "color": "949494"},
]

# ── CTA templates ───────────────────────────────────────────────
CTA_TEMPLATES = {
    0: {
        "M": "Jeśli tak, proszę o informację, jak będzie Panu wygodnie porozmawiać.",
        "F": "Jeśli tak, proszę o informację, jak będzie Pani wygodnie porozmawiać.",
    },
    1: {
        "M": "Jeśli widzi Pan sens takiej rozmowy, proszę o krótką odpowiedź - mogę dopasować się do telefonu albo Teams.",
        "F": "Jeśli widzi Pani sens takiej rozmowy, proszę o krótką odpowiedź - mogę dopasować się do telefonu albo Teams.",
    },
    2: {
        "M": "Proszę dać znać, czy wygodniejsza będzie krótka rozmowa telefoniczna czy Teams.",
        "F": "Proszę dać znać, czy wygodniejsza będzie krótka rozmowa telefoniczna czy Teams.",
    },
}


def sanitize_text(text: str) -> str:
    """Replace em dashes with hyphens per global rules."""
    return text.replace("\u2014", " - ")


def fix_role_phrasing(text: str) -> str:
    """Replace 'Z perspektywy [stanowisko]...' with 'W Pana/Pani roli jako [stanowisko]...'"""
    import re
    text = re.sub(
        r"Z perspektywy ([\w\s,/]+)",
        lambda m: f"W roli jako {m.group(1)}",
        text
    )
    return text


# ── Load context files ──────────────────────────────────────────
def load_context_files() -> dict:
    ctx_dir = os.path.join(_PROJECT_ROOT, "context")
    files = {}
    if not os.path.isdir(ctx_dir):
        return files
    for fname in sorted(os.listdir(ctx_dir)):
        if fname.endswith(".md"):
            fpath = os.path.join(ctx_dir, fname)
            with open(fpath, "r", encoding="utf-8") as f:
                files[fname] = f.read()
    return files


# ── LLM call ────────────────────────────────────────────────────
def generate_emails_for_contact(contact: dict, context_files: dict) -> dict | None:
    """Call LLM to generate 3-email sequence."""

    first_name = contact.get("first_name", "")
    vocative = VOCATIVE_FIX.get(first_name, contact.get("first_name_vocative", first_name))
    gender_str = contact.get("recipient_gender", "male")
    greeting = "Panie" if gender_str == "male" else "Pani"

    prompt_path = os.path.join(_ADHOC_ROOT, "prompts", "adhoc_email_writer.md")

    payload = {
        "contact_name": contact["name"],
        "contact_title": contact.get("apollo_title") or contact.get("known_title", ""),
        "company": contact.get("apollo_company") or contact.get("company", ""),
        "recipient_gender": gender_str,
        "first_name_vocative": vocative,
        "greeting_formal": greeting,
        "panel_title": contact.get("panel_title", ""),
        "post_url": contact.get("post_url", ""),
        "event_name": contact.get("event", "Poland & CEE Retail Summit 2026"),
        "key_themes": contact.get("key_themes", []),
        "business_problems": contact.get("business_problems", []),
        "personalization_angle": contact.get("personalization_angle", ""),
        "hypothesis": contact.get("negotiation_intelligence_angle", ""),
        "suggested_framework": "Cost/Trend/Dostawca/Decyzja",
        "industry_fit": contact.get("industry_fit", ""),
    }

    print(f"  [LLM] Calling for {contact['name']} (vocative: {vocative}, gender: {gender_str})...")

    result = generate_json(
        agent_name="AdHocEmailWriter",
        prompt_path=prompt_path,
        user_payload=payload,
        context_files=context_files,
        relevant_context_keys=["00_master", "01_offer", "03_messaging", "05_quality"],
        max_tokens=3000,
        temperature=0.5,
    )

    return result


# ── DOCX generation (v3 Outlook-ready format) ──────────────────

def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))


def add_run(paragraph, text, size=11, bold=False, color="000000", font="Aptos"):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=12, after_pt=12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_signature_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=11, after_pt=0)
    add_run(p, "Pozdrawiam serdecznie,", size=11, color="000000")

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before_pt=0, after_pt=0)

    for sig in SIG_LINES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=0)
        if sig["text"] == "":
            add_run(p, "", size=sig["size"])
        else:
            add_run(p, sig["text"], size=sig["size"], bold=sig["bold"], color=sig["color"])


def add_email_block(doc, email_num, subject, body, gender_code, email_idx):
    """Add one email block to the document (Outlook-ready format)."""

    body = sanitize_text(body)
    body = fix_role_phrasing(body)
    subject = sanitize_text(subject)

    # Reference label
    labels = ["EMAIL 1", "EMAIL 2", "EMAIL 3"]
    ref_p = doc.add_paragraph()
    set_paragraph_spacing(ref_p, before_pt=0, after_pt=4)
    add_run(ref_p, labels[email_num], size=10, bold=True, color="999999")

    # Subject line
    temat_p = doc.add_paragraph()
    set_paragraph_spacing(temat_p, before_pt=0, after_pt=8)
    add_run(temat_p, "Temat: ", size=9, bold=True, color="999999")
    add_run(temat_p, subject, size=9, bold=True, color="666666")

    # Body text
    body_clean = body.strip()
    # Remove trailing closing if LLM added it
    for closing in [
        "Pozdrawiam serdecznie,\nTomasz Uściński",
        "Pozdrawiam serdecznie,\n\nTomasz Uściński",
        "Pozdrawiam serdecznie,",
    ]:
        if body_clean.endswith(closing):
            body_clean = body_clean[: -len(closing)].rstrip()
            break

    paragraphs = body_clean.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=6)
        lines = para_text.split("\n")
        for i, line in enumerate(lines):
            add_run(p, line, size=11, color="000000")
            if i < len(lines) - 1:
                p.add_run("\n")

    # CTA line
    cta_text = CTA_TEMPLATES.get(email_idx, {}).get(gender_code, "")
    if cta_text:
        cta_p = doc.add_paragraph()
        set_paragraph_spacing(cta_p, before_pt=0, after_pt=6)
        add_run(cta_p, cta_text, size=11, color="000000")

    # Signature
    add_signature_block(doc)


def generate_v3_doc(contact: dict, emails_data: dict, output_dir: str) -> str:
    """Generate a single v3 Outlook-ready Word document."""

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    name = contact["name"]
    title = contact.get("apollo_title") or contact.get("known_title", "")
    company = contact.get("apollo_company") or contact.get("company", "")
    panel = sanitize_text(contact.get("panel_title", ""))
    gender_str = contact.get("recipient_gender", "male")
    gender_code = "M" if gender_str == "male" else "F"

    # File name from contact name
    nfkd = unicodedata.normalize("NFKD", name.lower().replace(" ", "_"))
    filename_base = "".join(c for c in nfkd if not unicodedata.combining(c))
    # Handle special chars
    filename_base = filename_base.replace("-", "_")

    # Document header
    header_p = doc.add_paragraph()
    set_paragraph_spacing(header_p, before_pt=0, after_pt=2)
    add_run(header_p, name, size=14, bold=True, color="333333")

    info_parts = []
    if title:
        info_parts.append(title)
    if company:
        info_parts.append(company)
    info_parts.append("(brak adresu e-mail)")
    info_p = doc.add_paragraph()
    set_paragraph_spacing(info_p, before_pt=0, after_pt=2)
    add_run(info_p, " | ".join(info_parts), size=9, color="888888")

    panel_p = doc.add_paragraph()
    set_paragraph_spacing(panel_p, before_pt=0, after_pt=0)
    add_run(panel_p, f"Panel: {panel}", size=9, color="AAAAAA")

    # Three email blocks
    for i in range(3):
        subject = emails_data.get(f"email_{i+1}_subject", "")
        body = emails_data.get(f"email_{i+1}_body", "")
        if not subject and not body:
            continue
        add_horizontal_line(doc)
        add_email_block(doc, email_num=i, subject=subject, body=body,
                        gender_code=gender_code, email_idx=i)

    filename = f"{filename_base}_v3.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


# ── Main ────────────────────────────────────────────────────────
def main():
    # Paths
    outputs_dir = os.path.join(_ADHOC_ROOT, "output")
    target_run = "2026-04-16_12-46-26_adhoc_linkedin_retail_summit_2026"
    run_dir = os.path.join(outputs_dir, target_run)
    docs_dir = os.path.join(run_dir, "DOCS")
    os.makedirs(docs_dir, exist_ok=True)

    # Load campaign results
    results_path = os.path.join(run_dir, "campaign_results.json")
    with open(results_path, "r", encoding="utf-8") as f:
        all_results = json.load(f)

    # Filter no-email contacts
    no_email_contacts = [r for r in all_results if not r.get("apollo_email")]
    print(f"Found {len(no_email_contacts)} contacts without email addresses.\n")

    # Check LLM availability
    if not is_llm_available():
        print("ERROR: LLM is not available. Check GITHUB_TOKEN and LLM_PROVIDER env vars.")
        print("Required: LLM_PROVIDER=github, GITHUB_TOKEN=<token>")
        sys.exit(1)

    # Load context files
    context_files = load_context_files()
    print(f"Loaded {len(context_files)} context files.\n")

    generated = []
    failed = []

    for contact in no_email_contacts:
        name = contact["name"]
        print(f"Processing: {name}")

        # Generate emails via LLM
        emails_data = generate_emails_for_contact(contact, context_files)

        if not emails_data or not emails_data.get("email_1_body"):
            print(f"  ✗ FAILED - no email content generated for {name}")
            failed.append(name)
            continue

        model_used = emails_data.get("_llm_model_used", "unknown")
        print(f"  ✓ Emails generated (model: {model_used})")

        # Generate v3 DOCX
        filepath = generate_v3_doc(contact, emails_data, docs_dir)
        print(f"  ✓ Saved: {os.path.basename(filepath)}")
        generated.append(filepath)

        # Also update the campaign_results.json record
        contact["email_1_subject"] = emails_data.get("email_1_subject", "")
        contact["email_1_body"] = emails_data.get("email_1_body", "")
        contact["email_2_subject"] = emails_data.get("email_2_subject", "")
        contact["email_2_body"] = emails_data.get("email_2_body", "")
        contact["email_3_subject"] = emails_data.get("email_3_subject", "")
        contact["email_3_body"] = emails_data.get("email_3_body", "")
        contact["_llm_model_used"] = model_used
        contact["status"] = "ready_for_review_no_email"

    # Save updated campaign_results.json
    with open(results_path, "w", encoding="utf-8") as f:
        json.dump(all_results, f, ensure_ascii=False, indent=2)
    print(f"\nUpdated campaign_results.json with email content.")

    print(f"\n{'='*60}")
    print(f"DONE. Generated: {len(generated)}, Failed: {len(failed)}")
    if failed:
        print(f"Failed contacts: {', '.join(failed)}")
    print(f"Output: {docs_dir}")


if __name__ == "__main__":
    main()
