#!/usr/bin/env python3
"""
Generate v2 Word documents with SOURCE OF TRUTH email sequences.
Proper Aptos formatting, structured signature block, correct capitalization.
"""

import json
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ADHOC_ROOT = os.path.dirname(_SCRIPT_DIR)

SENDER_EMAIL_ADDR = "tomasz.uscinski@profitia.pl"

# ---------- SIGNATURE SPEC ----------
SIG_LINES = [
    {"text": "Tomasz Uściński", "size": 11, "bold": True, "color": "000000"},
    {"text": "Senior Client Partner | Procurement Technology", "size": 11, "bold": False, "color": "000000"},
    {"text": "+48 787 417 293", "size": 11, "bold": False, "color": "000000"},
    {"text": "", "size": 11, "bold": False, "color": "000000"},  # empty
    {"text": "PROFITIA Management Consultants Mazurowski i Wspólnicy Spółka Jawna", "size": 11, "bold": True, "color": "000000"},
    {"text": "SpendGuru Data-driven Procurement | Certyfikowany Partner CIPS w Polsce | Doradztwo | Szkolenia | Analityka zakupowa", "size": 11, "bold": False, "color": "000000"},
    {"text": "02-715 Warszawa, Villa Metro, ul. Puławska 145, V p.", "size": 11, "bold": False, "color": "000000"},
    {"text": "", "size": 11, "bold": False, "color": "000000"},  # empty
    {"text": "Uwaga: Ten e-mail jest poufny i przeznaczony tylko dla adresata (-ów) tej wiadomości. Jeżeli nie jesteś adresatem niniejszej wiadomości, usuń oryginał wiadomości wraz z wszelkimi wydrukami (kopiami) i załącznikami.", "size": 9, "bold": False, "color": "949494"},
]

# ---------- GENDER LOOKUP ----------
GENDER_LOOKUP = {
    "monika_kolaszynska": "F",
    "mariusz_makowski": "M",
    "nils_swolkien": "M",
    "kiril_marinov": "M",
    "robert_stupak": "M",
    "agnieszka_kreciszewska": "F",
    "justyna_zasowska": "F",
    "michal_slawski": "M",
    "maciej_gawronski": "M",
    "tomasz_suffner": "M",
    "konrad_marchlewski": "M",
    "marek_marzec": "M",
    "daniel_pralat": "M",
}

# ---------- CTA TEMPLATES (practical response) ----------
CTA_TEMPLATES = {
    0: {  # Email 1
        "M": "Jeśli tak, proszę o informację, jak będzie Panu wygodnie porozmawiać.",
        "F": "Jeśli tak, proszę o informację, jak będzie Pani wygodnie porozmawiać.",
    },
    1: {  # Email 2
        "M": "Jeśli widzi Pan sens takiej rozmowy, proszę o krótką odpowiedź - mogę dopasować się do telefonu albo Teams.",
        "F": "Jeśli widzi Pani sens takiej rozmowy, proszę o krótką odpowiedź - mogę dopasować się do telefonu albo Teams.",
    },
    2: {  # Email 3
        "M": "Proszę dać znać, czy wygodniejsza będzie krótka rozmowa telefoniczna czy Teams.",
        "F": "Proszę dać znać, czy wygodniejsza będzie krótka rozmowa telefoniczna czy Teams.",
    },
}


def sanitize_text(text: str) -> str:
    """Replace em dashes with regular hyphens per global campaign rules."""
    return text.replace("\u2014", " - ")


# ---------- SOURCE OF TRUTH — 13 contacts ----------
CONTACTS_V2 = [
    {
        "filename": "monika_kolaszynska",
        "name": "Monika Kolaszyńska",
        "title": "Chief Executive Officer",
        "company": "Super-Pharm Poland",
        "panel": "Ambicja to za mało — Strategy & Growth",
        "emails": [
            {
                "subject": "Super-Pharm, skala wzrostu i spójność decyzji kosztowych",
                "body": "Dzień dobry Pani Moniko,\n\ndziękuję za bardzo interesujący głos w panelu \u201eAmbicja to za mało \u2014 Strategy & Growth\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pani roli jako CEO Super-Pharm Poland szczególnie trafne wydało mi się napięcie między ambicją wzrostu a zdolnością organizacji do utrzymania spójności w codziennych decyzjach. W sieci takiej jak Super-Pharm to nie dotyczy wyłącznie strategii czy ekspansji, ale też jakości decyzji kosztowych, zakupowych i dostawczych, które później wpływają na marżę, dostępność i przewidywalność budżetu.\n\nW Profitii pracujemy właśnie na tym styku: koszt, trend rynkowy, sytuacja dostawcy i decyzja biznesowa. To pomaga uporządkować rozmowy z dostawcami i budować większą przewidywalność w kategoriach, w których presja kosztowa i promocyjna jest wysoka.\n\nCzy znalazłaby Pani 15\u201320 minut na krótką rozmowę? Chętnie wymienię się perspektywą, jak taki model wspiera skalowanie biznesu bez utraty kontroli nad marżą.",
            },
            {
                "subject": "Re: Super-Pharm, skala wzrostu i spójność decyzji kosztowych",
                "body": "Dzień dobry Pani Moniko,\n\nwracam do mojego maila, bo w praktyce właśnie na poziomie sieci detalicznej najczęściej widać, że wzrost komplikuje się nie przez brak ambicji, ale przez rozjazd między strategią a codziennymi decyzjami operacyjnymi.\n\nW modelu takim jak Super-Pharm dotyczy to szczególnie kategorii, gdzie jednocześnie trzeba pilnować marży, dostępności, presji promocyjnej i relacji z dostawcami. Dobrze uporządkowany standard oceny kosztu, trendu i ryzyka dostawcy pomaga lepiej decydować, kiedy renegocjować warunki, kiedy bronić budżetu, a kiedy zwiększać elastyczność.\n\nJeśli temat jest dla Pani aktualny, chętnie pokażę to na prostym przykładzie jednej kategorii.",
            },
            {
                "subject": "Re: krótka wymiana perspektyw?",
                "body": "Dzień dobry Pani Moniko,\n\nzostawię już tylko krótkie domknięcie. Jeśli uzna Pani, że warto porozmawiać o tym, jak wzmocnić przewidywalność decyzji kosztowych i zakupowych w skali Super-Pharm, chętnie dopasuję się do Pani kalendarza.\n\nMoże to być też krótki sanity check dla jednej wybranej kategorii.",
            },
        ],
    },
    {
        "filename": "mariusz_makowski",
        "name": "Mariusz Makowski",
        "title": "",
        "company": "Mars Wrigley",
        "panel": "Ambicja to za mało — Strategy & Growth",
        "emails": [
            {
                "subject": "Mars Wrigley i spójność decyzji wspierających wzrost",
                "body": "Dzień dobry Panie Mariuszu,\n\ndziękuję za bardzo interesujący głos w panelu \u201eAmbicja to za mało \u2014 Strategy & Growth\u201d podczas Poland & CEE Retail Summit 2026.\n\nBardzo trafne wydało mi się Pana spojrzenie na to, że wzrost wymaga nie tylko ambicji, ale też spójności ludzi, procesów i decyzji. W organizacji takiej jak Mars Wrigley oznacza to również konieczność utrzymania dyscypliny kosztowej i zakupowej przy jednoczesnym wspieraniu skali, innowacji i rentowności.\n\nW Profitii często pracujemy właśnie z tym napięciem: jak uporządkować decyzje wokół kosztu, trendu rynkowego, ryzyka dostawcy i wpływu na marżę, tak aby zakupy realnie wspierały wzrost, a nie tylko reagowały na presję cenową.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie pokażę, jak taki model działa w praktyce na poziomie kategorii.",
            },
            {
                "subject": "Re: Mars Wrigley i spójność decyzji wspierających wzrost",
                "body": "Dzień dobry Panie Mariuszu,\n\nwracam do Pana, bo w firmach o dużej skali właśnie brak wspólnej logiki kosztowej i zakupowej często utrudnia realizację strategii wzrostu.\n\nW praktyce najczęściej widać to tam, gdzie presja na marżę, koszty surowców, opakowań czy logistyki zaczyna rozjeżdżać cele między biznesem, finansami i zakupami. Uporządkowanie tego przez jeden standard przygotowania decyzji i negocjacji daje znacznie większą przewidywalność.\n\nJeśli uzna Pan, że to ciekawy temat, chętnie wymienię się kilkoma praktycznymi obserwacjami.",
            },
            {
                "subject": "Re: krótka rozmowa o spójności decyzji?",
                "body": "Dzień dobry Panie Mariuszu,\n\nzostawię już tylko krótkie domknięcie. Jeśli temat spójności decyzji kosztowych i zakupowych w kontekście wzrostu jest dla Pana aktualny, chętnie porozmawiam 15 minut.\n\nMożemy też po prostu przejść przez jeden przykład kategorii, w której presja kosztowa najmocniej wpływa na marżę.",
            },
        ],
    },
    {
        "filename": "nils_swolkien",
        "name": "Nils Swolkien",
        "title": "CEO Castorama Poland, Kingfisher PLC Group Executive",
        "company": "Castorama Polska",
        "panel": "Łączenie systemów zwiększa nie tylko efektywność — ujawnia też miejsca podatne na ryzyka",
        "emails": [
            {
                "subject": "Castorama: integracja danych to także decyzje kosztowe",
                "body": "Dzień dobry Panie Nilsie,\n\ndziękuję za inspirujący głos w panelu \u201eŁączenie systemów zwiększa nie tylko efektywność \u2014 ujawnia też miejsca podatne na ryzyka\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pana roli jako CEO Castorama Polska i osoby działającej także na poziomie Group Executive bardzo trafne wydało mi się to, że integracja systemów nie rozwiązuje problemu sama z siebie. W modelu DIY retail pełna wartość pojawia się dopiero wtedy, gdy zintegrowane dane wspierają lepsze decyzje o dostępności, zapasie, kosztach i rozmowach z dostawcami \u2014 czyli tam, gdzie ostatecznie rozstrzyga się customer promise i marża.\n\nW Profitii pracujemy właśnie na tym styku: koszt, trend, dostawca, decyzja. Pomaga to lepiej oceniać moment negocjacji, ryzyko podwyżek i wpływ decyzji zakupowych na dostępność oraz wynik.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się perspektywą na ten obszar.",
            },
            {
                "subject": "Re: Castorama: integracja danych to także decyzje kosztowe",
                "body": "Dzień dobry Panie Nilsie,\n\nwracam, bo właśnie w dużych organizacjach retailowych najczęściej widać, że zintegrowany system nie daje jeszcze zintegrowanej decyzji.\n\nW Castoramie stawką nie jest przecież sama efektywność procesu, ale to, czy dane przekładają się na lepszą ocenę kosztu, właściwy moment negocjacji albo obronę marży w rozmowach z dostawcami. Jeśli nie, ryzyko bardzo szybko wraca w postaci braków, nadmiarów albo niepotrzebnie oddanej marży dostawcy.\n\nJeśli to dla Pana interesujące, chętnie pokażę prosty model, który pomaga ten obszar uporządkować.",
            },
            {
                "subject": "Re: krótka wymiana perspektyw?",
                "body": "Dzień dobry Panie Nilsie,\n\nzostawię krótkie domknięcie. Jeśli uzna Pan, że warto porozmawiać o tym, jak połączyć integrację danych z lepszą kontrolą kosztów i decyzji dostawczych w Castoramie, chętnie znajdę dogodny termin.\n\nMoże to być też szybki sanity check jednej kategorii lub obszaru ryzyka.",
            },
        ],
    },
    {
        "filename": "kiril_marinov",
        "name": "Kiril Marinov",
        "title": "General Manager & Board Member, Henkel Consumer Brands Poland",
        "company": "Henkel Consumer Brands Poland",
        "panel": "Ambicja to za mało — Strategy & Growth",
        "emails": [
            {
                "subject": "Henkel Consumer Brands: wzrost wymaga spójności kosztowej",
                "body": "Dzień dobry Panie Kirilu,\n\ndziękuję za bardzo interesujący głos w panelu \u201eAmbicja to za mało \u2014 Strategy & Growth\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pana roli jako General Manager i Board Member Henkel Consumer Brands Poland szczególnie mocno wybrzmiało dla mnie to, że wzrost i ambicja strategiczna muszą być podparte spójnym modelem działania. W FMCG ten warunek bardzo szybko schodzi do poziomu kosztów, warunków handlowych, przewidywalności kategorii i jakości decyzji zakupowych, które później bezpośrednio wpływają na marżę.\n\nW Profitii pomagamy uporządkować właśnie ten obszar: koszt, trend rynku, sytuacja dostawcy i decyzja negocjacyjna. Dzięki temu łatwiej utrzymać dyscyplinę kosztową bez utraty elastyczności biznesowej.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie podzielę się praktycznym spojrzeniem na ten temat.",
            },
            {
                "subject": "Re: Henkel Consumer Brands: wzrost wymaga spójności kosztowej",
                "body": "Dzień dobry Panie Kirilu,\n\nwracam, bo w firmach FMCG bardzo często widać, że strategia wzrostu zaczyna się komplikować tam, gdzie koszty kategorii i decyzje zakupowe przestają być zarządzane według jednej logiki.\n\nW praktyce chodzi o to, czy organizacja potrafi szybko ocenić, które wzrosty kosztów są realne, jak reagować na zmienność surowców i w których miejscach dostawca ma przestrzeń do ustępstw. To zwykle ma większy wpływ na marżę niż sama deklarowana ambicja wzrostu.\n\nJeśli uzna Pan, że to ciekawy temat, chętnie pokażę prosty przykład takiego podejścia.",
            },
            {
                "subject": "Re: 15 minut rozmowy?",
                "body": "Dzień dobry Panie Kirilu,\n\nzostawię krótkie domknięcie. Jeśli temat lepszej kontroli kosztów i przygotowania negocjacji w Henkel Consumer Brands Poland jest dla Pana aktualny, chętnie porozmawiam 15 minut.\n\nMożemy oprzeć rozmowę na jednej wybranej kategorii.",
            },
        ],
    },
    {
        "filename": "robert_stupak",
        "name": "Robert Stupak",
        "title": "Executive Director, Marketing, E-Commerce, Pricing, Tech & Data; Member of the Board",
        "company": "Carrefour",
        "panel": "Łączenie systemów zwiększa nie tylko efektywność — ujawnia też miejsca podatne na ryzyka",
        "emails": [
            {
                "subject": "Carrefour: dane są wartościowe dopiero, gdy poprawiają decyzję",
                "body": "Dzień dobry Panie Robercie,\n\ndziękuję za bardzo interesujący głos w panelu \u201eŁączenie systemów zwiększa nie tylko efektywność \u2014 ujawnia też miejsca podatne na ryzyka\u201d podczas Poland & CEE Retail Summit 2026.\n\nDla osoby odpowiadającej w Carrefour jednocześnie za marketing, e-commerce, pricing, tech i data szczególnie trafne było to, że integracja danych nie jest celem samym w sobie. Jej realna wartość pojawia się dopiero wtedy, gdy poprawia decyzje handlowe i operacyjne \u2014 od ceny i promocji, przez dostępność, po warunki zakupowe i marżę.\n\nW Profitii pracujemy właśnie na takim styku. Łączymy analizę kosztu, trendów i ryzyka dostawcy z decyzją biznesową, tak aby dane nie kończyły się w dashboardzie, tylko wzmacniały negocjacje i przewidywalność kategorii.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się perspektywą na ten obszar.",
            },
            {
                "subject": "Re: Carrefour: dane są wartościowe dopiero, gdy poprawiają decyzję",
                "body": "Dzień dobry Panie Robercie,\n\nwracam, bo przy tak szerokiej odpowiedzialności jak pricing, e-commerce i data bardzo szybko widać, czy organizacja naprawdę pracuje na wspólnym obrazie kategorii.\n\nW praktyce największy problem pojawia się wtedy, gdy dane są zintegrowane technologicznie, ale nadal nie przekładają się na lepszą ocenę kosztu, właściwy moment negocjacji albo obronę marży w rozmowach z dostawcami. Wtedy ryzyko wraca inną drogą \u2014 przez pricing, dostępność albo rentowność promocji.\n\nJeśli to dla Pana interesujące, chętnie pokażę, jak firmy porządkują ten obszar w praktyce.",
            },
            {
                "subject": "Re: krótka rozmowa o danych i marży?",
                "body": "Dzień dobry Panie Robercie,\n\nzostawię już krótkie domknięcie. Jeśli uzna Pan, że warto porozmawiać o tym, jak połączyć dane, pricing i decyzje zakupowe w sposób bardziej wspierający marżę Carrefour, chętnie znajdę dogodny termin.\n\nMoże to być też krótki sanity check dla jednej kategorii.",
            },
        ],
    },
    {
        "filename": "agnieszka_kreciszewska",
        "name": "Agnieszka Kreciszewska",
        "title": "",
        "company": "Caparol Polska",
        "panel": "Financial Performance & Business Management — marka własna jako gra o model biznesowy",
        "emails": [
            {
                "subject": "Marka własna: gdzie naprawdę ucieka marża",
                "body": "Dzień dobry Pani Agnieszko,\n\ndziękuję za bardzo interesujący głos w panelu \u201eFinancial Performance & Business Management \u2014 marka własna jako gra o model biznesowy\u201d podczas Poland & CEE Retail Summit 2026.\n\nBardzo trafne wydało mi się to, że w markach własnych sprzedaż i wzrost wolumenu nie muszą automatycznie oznaczać poprawy wyniku. W firmie takiej jak Caparol Polska stawką jest przecież nie tylko koszt wytworzenia, ale też jakość decyzji o dostawcach, strukturze kosztu, presji sieci i przestrzeni do obrony marży.\n\nW Profitii pomagamy uporządkować ten obszar przez prostą logikę: koszt, trend, dostawca, decyzja. Dzięki temu łatwiej ocenić, gdzie oferta lub koszt są faktycznie uzasadnione, a gdzie warto wrócić do rozmowy z partnerem lub dostawcą.\n\nCzy znalazłaby Pani 15\u201320 minut na krótką rozmowę? Chętnie wymienię się praktycznymi obserwacjami z tego obszaru.",
            },
            {
                "subject": "Re: Marka własna: gdzie naprawdę ucieka marża",
                "body": "Dzień dobry Pani Agnieszko,\n\nwracam, bo w modelu private label najczęściej problem nie leży w jednym wskaźniku, tylko w tym, że koszt, warunki handlowe i ryzyko dostawcy są oceniane osobno.\n\nW praktyce prowadzi to do sytuacji, w której firma widzi sprzedaż, ale nie zawsze ma równie dobrą kontrolę nad tym, gdzie realnie oddaje marżę: w koszcie komponentu, w presji sieci, w konstrukcji oferty albo w zbyt słabej pozycji negocjacyjnej. Dobrze uporządkowany standard oceny dostawcy i kosztu potrafi bardzo pomóc.\n\nJeśli temat jest dla Pani aktualny, chętnie pokażę to na prostym przykładzie.",
            },
            {
                "subject": "Re: krótka wymiana perspektyw?",
                "body": "Dzień dobry Pani Agnieszko,\n\nzostawię krótkie domknięcie. Jeśli temat kontroli kosztu i ochrony marży w modelu private label jest dla Pani aktualny, chętnie porozmawiam 15 minut.\n\nMożemy oprzeć rozmowę na jednym wybranym obszarze kosztowym lub dostawcy.",
            },
        ],
    },
    {
        "filename": "justyna_zasowska",
        "name": "Justyna Zasowska",
        "title": "CFO, OBI Polska",
        "company": "OBI Polska",
        "panel": "Płynność, cash flow i kapitał obrotowy — rola decyzji zakupowych i predykcji",
        "emails": [
            {
                "subject": "OBI Polska: cash flow zaczyna się wcześniej niż w finansach",
                "body": "Dzień dobry Pani Justyno,\n\ndziękuję za bardzo interesujący głos w panelu \u201ePłynność, cash flow i kapitał obrotowy \u2014 rola decyzji zakupowych i predykcji\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pani roli jako CFO OBI Polska szczególnie trafne wydało mi się to, że cash flow jest w dużej mierze efektem jakości wcześniejszych decyzji zakupowych. W modelu DIY retail, przy dużej liczbie SKU i sezonowości, nawet drobne błędy w timingu zakupu, poziomie zapasu czy ocenie ryzyka dostawcy bardzo szybko przekładają się na gotówkę zamrożoną w zapasach i presję na marżę.\n\nW Profitii pomagamy porządkować ten obszar przez jedną logikę: koszt, trend, dostawca, decyzja. Dzięki temu łatwiej ocenić, kiedy kupować, jak negocjować i jak ograniczać ryzyko nadmiaru lub braku zapasu.\n\nCzy znalazłaby Pani 15\u201320 minut na krótką rozmowę? Chętnie wymienię się praktyczną perspektywą.",
            },
            {
                "subject": "Re: OBI Polska: cash flow zaczyna się wcześniej niż w finansach",
                "body": "Dzień dobry Pani Justyno,\n\nwracam, bo właśnie w roli CFO bardzo dobrze widać, że kapitał obrotowy jest późnym skutkiem wcześniejszych decyzji kosztowych i zakupowych.\n\nW praktyce chodzi o to, czy organizacja potrafi połączyć prognozę, koszt kategorii, warunki dostawcy i wpływ na gotówkę w jedną decyzję. Bez tego łatwo o nadmiar zapasu, zły timing kontraktacji albo zaakceptowanie warunków, które niepotrzebnie obciążają cash flow.\n\nJeśli uzna Pani, że temat jest wart krótkiej rozmowy, chętnie pokażę prosty przykład takiego podejścia.",
            },
            {
                "subject": "Re: 15 minut o kapitale obrotowym?",
                "body": "Dzień dobry Pani Justyno,\n\nzostawię krótkie domknięcie. Jeśli uzna Pani, że warto porozmawiać o tym, jak lepiej połączyć decyzje zakupowe z przewidywalnością kapitału obrotowego w OBI Polska, chętnie znajdę dogodny termin.\n\nMoże to być też szybki sanity check jednej kategorii.",
            },
        ],
    },
    {
        "filename": "michal_slawski",
        "name": "Michał Śławski",
        "title": "CFO, Castorama Polska",
        "company": "Castorama Polska",
        "panel": "Płynność, cash flow i kapitał obrotowy — rola decyzji zakupowych i predykcji",
        "emails": [
            {
                "subject": "Castorama: zapas, gotówka i decyzje zakupowe",
                "body": "Dzień dobry Panie Michale,\n\ndziękuję za bardzo interesujący głos w panelu \u201ePłynność, cash flow i kapitał obrotowy \u2014 rola decyzji zakupowych i predykcji\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pana roli jako CFO Castorama Polska bardzo trafne było dla mnie to, że przy dużej liczbie SKU płynność i gotówka są bezpośrednim skutkiem jakości decyzji zakupowych. W takim modelu biznesowym stawką nie jest przecież wyłącznie poziom zapasu, ale też timing kontraktacji, ekspozycja na zmienność kosztów i to, ile gotówki organizacja oddaje przez nieprecyzyjne decyzje.\n\nW Profitii pomagamy porządkować te decyzje przez prosty model: koszt, trend rynku, ryzyko dostawcy i decyzja biznesowa. To zwykle poprawia przewidywalność cash flow i ogranicza niepotrzebne napięcie między dostępnością a kapitałem obrotowym.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się praktycznym spojrzeniem na ten obszar.",
            },
            {
                "subject": "Re: Castorama: zapas, gotówka i decyzje zakupowe",
                "body": "Dzień dobry Panie Michale,\n\nwracam, bo w roli CFO najlepiej widać, że gotówka zamrożona w zapasie to zwykle nie problem magazynu, tylko wcześniejszej decyzji.\n\nW praktyce chodzi o to, czy firma potrafi odpowiednio wcześnie połączyć prognozę popytu, koszt kategorii, warunki dostawcy i wpływ na kapitał obrotowy. Bez tej spójności łatwo o nadmiary, braki albo decyzje zakupowe, które wyglądają dobrze lokalnie, ale osłabiają wynik finansowy całości.\n\nJeśli to dla Pana interesujące, chętnie pokażę prosty sposób uporządkowania tego obszaru.",
            },
            {
                "subject": "Re: krótka rozmowa o cash flow?",
                "body": "Dzień dobry Panie Michale,\n\nzostawię już krótkie domknięcie. Jeśli uzna Pan, że warto porozmawiać o tym, jak lepiej połączyć decyzje zakupowe z kontrolą zapasu i przewidywalnością gotówki w Castoramie, chętnie dopasuję się do Pana kalendarza.",
            },
        ],
    },
    {
        "filename": "maciej_gawronski",
        "name": "Maciej Gawroński",
        "title": "CEO, Idkfa Idclip / part of Fozzy Group",
        "company": "Idkfa Idclip / Fozzy Group",
        "panel": "Supply Chain Plenary Stage — zmienność jako codzienność",
        "emails": [
            {
                "subject": "Zmienność w supply chain a odporność wyniku",
                "body": "Dzień dobry Panie Macieju,\n\ndziękuję za bardzo interesujący głos w panelu \u201eSupply Chain Plenary Stage \u2014 zmienność jako codzienność\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pana roli jako CEO bardzo trafne wydało mi się to, że zmienność nie jest dziś już wyjątkiem, tylko warunkiem działania. W grupach retailowo-dystrybucyjnych oznacza to, że odporność biznesu buduje się nie tylko logistyką, ale też jakością decyzji kosztowych, dostawczych i negocjacyjnych \u2014 bo to one później wpływają na marżę, dostępność i przewidywalność.\n\nW Profitii pomagamy porządkować ten obszar przez prostą logikę: koszt, trend, dostawca, decyzja. Dzięki temu łatwiej ocenić, kiedy zabezpieczać warunki, gdzie dostawca niesie największe ryzyko i jak reagować na zmienność bez nadmiernego chaosu.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się perspektywą na ten obszar.",
            },
            {
                "subject": "Re: Zmienność w supply chain a odporność wyniku",
                "body": "Dzień dobry Panie Macieju,\n\nwracam, bo właśnie w warunkach trwałej zmienności najwięcej wartości daje nie sama szybkość reakcji, ale jakość logiki decyzyjnej.\n\nW praktyce firmy najczęściej tracą tam, gdzie nie łączą w jednej decyzji kosztu, prognozy, ryzyka dostawcy i wpływu na wynik. Wtedy reakcja jest szybka, ale nie zawsze optymalna. Dobrze uporządkowany model pomaga ograniczyć ten efekt i lepiej przygotować rozmowy z partnerami.\n\nJeśli uzna Pan, że warto, chętnie pokażę prosty przykład takiego podejścia.",
            },
            {
                "subject": "Re: 15 minut o zarządzaniu zmiennością?",
                "body": "Dzień dobry Panie Macieju,\n\nzostawię krótkie domknięcie. Jeśli temat lepszego połączenia zmienności supply chain z decyzjami kosztowymi i negocjacyjnymi jest dla Pana aktualny, chętnie porozmawiam 15 minut.",
            },
        ],
    },
    {
        "filename": "tomasz_suffner",
        "name": "Tomasz Suffner",
        "title": "General Manager, Unilever Personal Care Poland; PC Europe Leadership Team Member",
        "company": "Unilever Personal Care Poland",
        "panel": "Zmiana stała się normą. Stabilność – niekoniecznie.",
        "emails": [
            {
                "subject": "Unilever: stabilność nie z rynku, tylko z jakości decyzji",
                "body": "Dzień dobry Panie Tomaszu,\n\ndziękuję za interesujący głos w panelu \u201eZmiana stała się normą. Stabilność \u2013 niekoniecznie.\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pana roli jako General Manager Unilever Personal Care Poland bardzo trafne wydało mi się to, że w zmiennym otoczeniu stabilność nie wynika już z samego rynku, tylko z jakości procesów i decyzji wewnątrz organizacji. W FMCG oznacza to także zdolność do konsekwentnego zarządzania kosztami kategorii, zmianą surowców, warunkami dostawców i wpływem tych decyzji na marżę.\n\nW Profitii pomagamy budować właśnie taki bardziej powtarzalny standard: koszt, trend, dostawca, decyzja. To porządkuje moment negocjacji, ocenę zasadności podwyżek i sposób obrony budżetu.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się praktycznym spojrzeniem na ten obszar.",
            },
            {
                "subject": "Re: Unilever: stabilność nie z rynku, tylko z jakości decyzji",
                "body": "Dzień dobry Panie Tomaszu,\n\nwracam, bo w praktyce wiele organizacji szuka dziś stabilności w otoczeniu, którego już nie da się ustabilizować.\n\nZnacznie lepszy efekt daje uporządkowanie tego, jak firma ocenia koszt, trend rynku, ryzyko partnera i moment decyzji. Wtedy nawet przy zmienności łatwiej utrzymać spójność działań i lepiej bronić marży w kluczowych kategoriach.\n\nJeśli uzna Pan, że to ciekawy temat, chętnie pokażę prosty przykład takiego podejścia.",
            },
            {
                "subject": "Re: krótka wymiana perspektyw?",
                "body": "Dzień dobry Panie Tomaszu,\n\nzostawię krótkie domknięcie. Jeśli uzna Pan, że warto porozmawiać o tym, jak zwiększyć powtarzalność decyzji kosztowych i negocjacyjnych w zmiennym rynku FMCG, chętnie znajdę dogodny termin.",
            },
        ],
    },
    {
        "filename": "konrad_marchlewski",
        "name": "Konrad Marchlewski",
        "title": "",
        "company": "BEWA",
        "panel": "Ujednolicanie jakości danych, pomiaru retail media i kontroli personalizacji w celu poprawy decyzji finansowych",
        "emails": [
            {
                "subject": "BEWA: dane są użyteczne dopiero, gdy poprawiają rentowność decyzji",
                "body": "Dzień dobry Panie Konradzie,\n\ndziękuję za cenny głos w panelu \u201eUjednolicanie jakości danych, pomiaru retail media i kontroli personalizacji w celu poprawy decyzji finansowych\u201d podczas Poland & CEE Retail Summit 2026.\n\nBardzo trafne wydało mi się to, że jakość danych nie jest dziś tematem technicznym, tylko finansowym i zarządczym. W firmie takiej jak BEWA problem zwykle nie polega na samym braku danych, ale na tym, czy dają one spójny obraz rentowności, kosztów i jakości decyzji handlowo-zakupowych.\n\nW Profitii pracujemy właśnie na tym styku: koszt, trend, dostawca, decyzja. To pomaga uporządkować ocenę kosztów, ryzyka i przestrzeni negocjacyjnej tak, aby dane realnie wspierały marżę, a nie tylko raportowanie.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się praktyczną perspektywą.",
            },
            {
                "subject": "Re: BEWA: dane są użyteczne dopiero, gdy poprawiają rentowność decyzji",
                "body": "Dzień dobry Panie Konradzie,\n\nwracam, bo w wielu firmach problem nie polega już na tym, czy dane są dostępne, tylko czy prowadzą do jednej, sensownej decyzji.\n\nJeśli koszt, sprzedaż, rentowność i ryzyko dostawcy są interpretowane osobno, bardzo łatwo o decyzje, które lokalnie wyglądają dobrze, ale globalnie osłabiają marżę albo pozycję negocjacyjną. Dobrze uporządkowany model oceny potrafi to bardzo szybko poprawić.\n\nJeśli uzna Pan, że warto, chętnie pokażę prosty przykład.",
            },
            {
                "subject": "Re: 15 minut rozmowy?",
                "body": "Dzień dobry Panie Konradzie,\n\nzostawię krótkie domknięcie. Jeśli temat lepszego połączenia jakości danych z rentownością decyzji w BEWA jest dla Pana aktualny, chętnie porozmawiam 15 minut.",
            },
        ],
    },
    {
        "filename": "marek_marzec",
        "name": "Marek Marzec",
        "title": "",
        "company": "PPH Ewa Bis",
        "panel": "Określanie priorytetów przywódczych równoważących oczekiwania właścicieli, ambicje strategiczne i realia ekonomiczne",
        "emails": [
            {
                "subject": "Strategia w produkcji zaczyna się od decyzji kosztowych",
                "body": "Dzień dobry Panie Marku,\n\ndziękuję za bardzo interesujący głos w panelu \u201eOkreślanie priorytetów przywódczych równoważących oczekiwania właścicieli, ambicje strategiczne i realia ekonomiczne\u201d podczas Poland & CEE Retail Summit 2026.\n\nBardzo trafne wydało mi się to napięcie między ambicją właścicielską a codzienną rzeczywistością firmy. W organizacji produkcyjnej takiej jak PPH Ewa Bis to zwykle nie rozstrzyga się na poziomie samej strategii, tylko na poziomie decyzji o koszcie surowców, warunkach zakupowych, przewidywalności kategorii i ochronie marży.\n\nW Profitii pomagamy porządkować właśnie ten obszar: koszt, trend, dostawca, decyzja. Dzięki temu łatwiej przełożyć cele strategiczne na bardziej spójne decyzje operacyjne i negocjacyjne.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się praktycznym spojrzeniem na ten temat.",
            },
            {
                "subject": "Re: Strategia w produkcji zaczyna się od decyzji kosztowych",
                "body": "Dzień dobry Panie Marku,\n\nwracam, bo w firmach produkcyjnych strategia bardzo często przegrywa nie przez brak ambicji, tylko przez zbyt słabą kontrolę codziennych decyzji kosztowych.\n\nJeśli organizacja nie ma jednej logiki oceny kosztu, trendu rynku i ryzyka dostawcy, wtedy łatwo o rozjazd między oczekiwaniem właścicieli a realiami operacyjnymi. To zwykle najmocniej uderza w marżę i przewidywalność wyniku.\n\nJeśli uzna Pan, że temat jest wart krótkiej rozmowy, chętnie pokażę prosty przykład takiego podejścia.",
            },
            {
                "subject": "Re: krótka wymiana perspektyw?",
                "body": "Dzień dobry Panie Marku,\n\nzostawię krótkie domknięcie. Jeśli warto porozmawiać o tym, jak lepiej połączyć priorytety strategiczne z decyzjami kosztowymi i zakupowymi w PPH Ewa Bis, chętnie znajdę dogodny termin.",
            },
        ],
    },
    {
        "filename": "daniel_pralat",
        "name": "Daniel Prałat",
        "title": "Group Commercial Director",
        "company": "CHATA POLSKA",
        "panel": "Określanie priorytetów przywódczych równoważących oczekiwania właścicieli, ambicje strategiczne i realia ekonomiczne",
        "emails": [
            {
                "subject": "CHATA POLSKA: marża handlowa potrzebuje wspólnej logiki danych i kosztu",
                "body": "Dzień dobry Panie Danielu,\n\ndziękuję za bardzo interesujący głos w panelu \u201eOkreślanie priorytetów przywódczych równoważących oczekiwania właścicieli, ambicje strategiczne i realia ekonomiczne\u201d podczas Poland & CEE Retail Summit 2026.\n\nW Pana roli jako Group Commercial Director CHATA POLSKA bardzo trafne wydało mi się to, że cele strategiczne często rozjeżdżają się na styku comercial, finansów i operacji. W organizacji handlowej takiej jak CHATA POLSKA widać to szczególnie tam, gdzie koszt zakupu, warunki handlowe i decyzje komercyjne nie są oceniane według jednej logiki \u2014 a to bardzo szybko przekłada się na marżę.\n\nW Profitii pomagamy porządkować ten obszar przez prosty model: koszt, trend, dostawca, decyzja. Dzięki temu łatwiej podejmować spójniejsze decyzje handlowo-zakupowe i lepiej przygotowywać rozmowy z dostawcami.\n\nCzy znalazłby Pan 15\u201320 minut na krótką rozmowę? Chętnie wymienię się perspektywą na ten temat.",
            },
            {
                "subject": "Re: CHATA POLSKA: marża handlowa potrzebuje wspólnej logiki danych i kosztu",
                "body": "Dzień dobry Panie Danielu,\n\nwracam, bo właśnie na poziomie comercial najczęściej widać, czy organizacja naprawdę pracuje na wspólnych przesłankach.\n\nJeśli koszt zakupu, trend rynku, ryzyko dostawcy i decyzja promocyjno-handlowa są rozdzielone, wtedy łatwo o ruchy, które wspierają sprzedaż krótkoterminowo, ale oddają marżę. Dobrze uporządkowany standard oceny i negocjacji pomaga ten efekt mocno ograniczyć.\n\nJeśli uzna Pan, że to ciekawy temat, chętnie pokażę prosty przykład.",
            },
            {
                "subject": "Re: 15 minut rozmowy?",
                "body": "Dzień dobry Panie Danielu,\n\nzostawię krótkie domknięcie. Jeśli temat lepszego połączenia decyzji handlowych, zakupowych i marżowych w CHATA POLSKA jest dla Pana aktualny, chętnie porozmawiam 15 minut.",
            },
        ],
    },
]


# ---------- Email address lookup from Apollo data ----------
def load_email_lookup(run_dir: str) -> dict:
    """Load campaign_results.json and build name → email lookup."""
    results_path = os.path.join(run_dir, "campaign_results.json")
    if not os.path.exists(results_path):
        return {}
    with open(results_path, "r", encoding="utf-8") as f:
        results = json.load(f)
    lookup = {}
    for r in results:
        name = r.get("name", "")
        email = r.get("apollo_email", "")
        if name and email:
            # normalize key: lowercase, strip diacritics-like
            key = name.lower().replace(" ", "_")
            lookup[key] = email
            # also store ascii-ized version
            import unicodedata
            nfkd = unicodedata.normalize("NFKD", key)
            ascii_key = "".join(c for c in nfkd if not unicodedata.combining(c))
            lookup[ascii_key] = email
    return lookup


# ---------- DOCX formatting ----------

def add_sig_run(paragraph, text, size_pt, bold, color_hex, font_name="Aptos"):
    """Add a formatted run to a paragraph for signature lines."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_formatted_signature(doc):
    """Add the formatted signature block to the document."""
    # Closing line
    closing_p = doc.add_paragraph()
    closing_p.space_before = Pt(8)
    closing_p.space_after = Pt(4)
    closing_run = closing_p.add_run("Pozdrawiam serdecznie,")
    closing_run.font.name = "Aptos"
    closing_run.font.size = Pt(11)
    closing_run.font.color.rgb = RGBColor(0, 0, 0)

    for sig in SIG_LINES:
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        if sig["text"] == "":
            # Empty line — just add a small run to preserve spacing
            run = p.add_run("")
            run.font.size = Pt(sig["size"])
        else:
            add_sig_run(p, sig["text"], sig["size"], sig["bold"], sig["color"])


def add_email_block(doc, step_label, subject, body, recipient_email, is_reply=False, email_idx=0, gender="M"):
    """Add one email block to the document."""

    # Global rule: replace em dashes
    body = sanitize_text(body)
    subject = sanitize_text(subject)
    step_label = sanitize_text(step_label)

    # Separator
    if is_reply:
        sep = doc.add_paragraph()
        sep_run = sep.add_run("\u2500" * 60)
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        sep.space_before = Pt(16)
        sep.space_after = Pt(4)

    # Step label
    label_p = doc.add_paragraph()
    label_run = label_p.add_run(step_label)
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    label_run.font.name = "Aptos"
    label_p.space_after = Pt(2)

    # Header lines (Od / Do / Temat)
    def add_header_line(label_text, value_text):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        lbl = p.add_run(f"{label_text}: ")
        lbl.bold = True
        lbl.font.size = Pt(9)
        lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        lbl.font.name = "Aptos"
        val = p.add_run(value_text)
        val.font.size = Pt(9)
        val.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        val.font.name = "Aptos"

    sender_str = f"Tomasz Uściński <tomasz.uscinski@profitia.pl>"
    add_header_line("Od", sender_str)
    add_header_line("Do", recipient_email)
    add_header_line("Temat", subject)

    # Spacer
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(4)
    spacer.space_after = Pt(4)

    # Body paragraphs
    body_clean = body.strip()
    # Remove trailing closing if present
    for closing in [
        "Pozdrawiam serdecznie,\nTomasz Uściński",
        "Pozdrawiam serdecznie,\n\nTomasz Uściński",
        "Podpis",
    ]:
        if body_clean.endswith(closing):
            body_clean = body_clean[: -len(closing)].rstrip()
            break

    paragraphs = body_clean.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        # Handle multi-line within a paragraph (single \n)
        lines = para_text.split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = "Aptos"
            run.font.color.rgb = RGBColor(0, 0, 0)
            if i < len(lines) - 1:
                p.add_run("\n").font.size = Pt(11)
        p.space_after = Pt(4)

    # Add practical CTA (global campaign rule)
    cta_text = CTA_TEMPLATES.get(email_idx, {}).get(gender, "")
    if cta_text:
        cta_p = doc.add_paragraph()
        cta_run = cta_p.add_run(cta_text)
        cta_run.font.size = Pt(11)
        cta_run.font.name = "Aptos"
        cta_run.font.color.rgb = RGBColor(0, 0, 0)
        cta_p.space_after = Pt(4)

    # Add formatted signature
    add_formatted_signature(doc)


def generate_v2_doc(contact: dict, email_addr: str, output_dir: str) -> str:
    """Generate a single v2 Word document."""

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Aptos"
    font.size = Pt(11)

    name = contact["name"]
    title = contact["title"]
    company = contact["company"]
    panel = contact["panel"]
    emails = contact["emails"]

    # Title
    title_p = doc.add_heading(level=1)
    title_run = title_p.add_run(name)
    title_run.font.size = Pt(16)
    title_run.font.name = "Aptos"

    # Contact info line
    info_p = doc.add_paragraph()
    info_parts = []
    if title:
        info_parts.append(title)
    if company:
        info_parts.append(company)
    if email_addr:
        info_parts.append(email_addr)
    info_run = info_p.add_run(" | ".join(info_parts))
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    info_run.font.name = "Aptos"

    # Panel
    if panel:
        panel_p = doc.add_paragraph()
        panel_run = panel_p.add_run(f"Panel: {sanitize_text(panel)}")
        panel_run.font.size = Pt(9)
        panel_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        panel_run.italic = True
        panel_run.font.name = "Aptos"

    # Thin separator
    doc.add_paragraph().space_after = Pt(2)

    step_labels = [
        "EMAIL 1 - Pierwszy kontakt",
        "EMAIL 2 - Follow-up",
        "EMAIL 3 - Final follow-up",
    ]

    gender = GENDER_LOOKUP.get(contact["filename"], "M")

    for i, email_data in enumerate(emails):
        add_email_block(
            doc,
            step_label=step_labels[i],
            subject=email_data["subject"],
            body=email_data["body"],
            recipient_email=email_addr,
            is_reply=(i > 0),
            email_idx=i,
            gender=gender,
        )

    filename = f"{contact['filename']}_v2.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def main():
    # Find run directory
    outputs_dir = os.path.join(_ADHOC_ROOT, "outputs")
    target_run = "2026-04-16_12-46-26_adhoc_linkedin_retail_summit_2026"
    run_dir = os.path.join(outputs_dir, target_run)
    docs_dir = os.path.join(run_dir, "DOCS")
    os.makedirs(docs_dir, exist_ok=True)

    # Load email lookup from Apollo data
    email_lookup = load_email_lookup(run_dir)
    print(f"Loaded {len(email_lookup)} email mappings from Apollo data.")

    generated = []
    for contact in CONTACTS_V2:
        # Try to find email address
        fn = contact["filename"]
        email_addr = email_lookup.get(fn, "")
        if not email_addr:
            # try other key forms
            name_key = contact["name"].lower().replace(" ", "_")
            import unicodedata
            nfkd = unicodedata.normalize("NFKD", name_key)
            ascii_key = "".join(c for c in nfkd if not unicodedata.combining(c))
            email_addr = email_lookup.get(ascii_key, email_lookup.get(name_key, ""))

        filepath = generate_v2_doc(contact, email_addr, docs_dir)
        print(f"  \u2713 {os.path.basename(filepath)}  ({email_addr or 'NO EMAIL'})")
        generated.append(filepath)

    print(f"\nDone. {len(generated)} v2 docs saved to:\n  {docs_dir}")


if __name__ == "__main__":
    main()
