#!/usr/bin/env python3
"""
Generate v3 Word documents - FINAL Outlook-ready format.
Each doc contains 3 emails ready to copy-paste into Outlook.

Format per email:
- Reference label (EMAIL 1/2/3) - small gray, not part of paste
- Do: recipient email - reference only
- Temat: subject line - copy to Outlook Subject field
- Body text - paste directly into Outlook body
- Signature - formatted per spec

All text in Aptos 11pt (Outlook default), proper spacing.
"""

import json
import os
import sys
import unicodedata
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_ADHOC_ROOT = os.path.dirname(_SCRIPT_DIR)

# Import contact data and constants from v2
sys.path.insert(0, _SCRIPT_DIR)
from generate_word_docs_v2 import (
    CONTACTS_V2,
    SIG_LINES,
    GENDER_LOOKUP,
    CTA_TEMPLATES,
    sanitize_text,
    load_email_lookup,
)


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing_pt=None):
    """Set exact paragraph spacing."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_spacing_pt:
        spacing.set(qn("w:line"), str(int(line_spacing_pt * 20)))
        spacing.set(qn("w:lineRule"), "exact")


def add_run(paragraph, text, size=11, bold=False, color="000000", font="Aptos"):
    """Add a formatted run."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_horizontal_line(doc):
    """Add a thin horizontal line as separator."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=12, after_pt=12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_signature_block(doc):
    """Add the full signature block per spec."""
    # "Pozdrawiam serdecznie," line
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=11, after_pt=0)
    add_run(p, "Pozdrawiam serdecznie,", size=11, color="000000")

    # Empty line between closing and signature
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before_pt=0, after_pt=0)

    for sig in SIG_LINES:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=0)
        if sig["text"] == "":
            add_run(p, "", size=sig["size"])
        else:
            add_run(p, sig["text"], size=sig["size"], bold=sig["bold"], color=sig["color"])


def add_email_block(doc, email_num, subject, body, recipient_email, gender, email_idx):
    """Add one complete email block to the document."""

    # Sanitize
    body = sanitize_text(body)
    subject = sanitize_text(subject)

    # ---- Reference header (gray, small - not part of paste) ----
    labels = ["EMAIL 1", "EMAIL 2", "EMAIL 3"]
    ref_p = doc.add_paragraph()
    set_paragraph_spacing(ref_p, before_pt=0, after_pt=4)
    add_run(ref_p, labels[email_num], size=10, bold=True, color="999999")

    # Do: line (reference)
    do_p = doc.add_paragraph()
    set_paragraph_spacing(do_p, before_pt=0, after_pt=0)
    add_run(do_p, "Do: ", size=9, bold=True, color="999999")
    add_run(do_p, recipient_email, size=9, color="999999")

    # Temat: line (user copies this to Outlook Subject field)
    temat_p = doc.add_paragraph()
    set_paragraph_spacing(temat_p, before_pt=0, after_pt=8)
    add_run(temat_p, "Temat: ", size=9, bold=True, color="999999")
    add_run(temat_p, subject, size=9, bold=True, color="666666")

    # ---- Body text (this is what gets pasted into Outlook) ----
    body_clean = body.strip()
    # Remove trailing closing if present
    for closing in [
        "Pozdrawiam serdecznie,\nTomasz Uściński",
        "Pozdrawiam serdecznie,\n\nTomasz Uściński",
        "Podpis",
    ]:
        if body_clean.endswith(closing):
            body_clean = body_clean[: -len(closing)].rstrip()
            break

    paragraphs = body_clean.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=6)
        lines = para_text.split("\n")
        for i, line in enumerate(lines):
            add_run(p, line, size=11, color="000000")
            if i < len(lines) - 1:
                p.add_run("\n")

    # CTA line
    cta_text = CTA_TEMPLATES.get(email_idx, {}).get(gender, "")
    if cta_text:
        cta_p = doc.add_paragraph()
        set_paragraph_spacing(cta_p, before_pt=0, after_pt=6)
        add_run(cta_p, cta_text, size=11, color="000000")

    # Signature
    add_signature_block(doc)


def generate_v3_doc(contact: dict, email_addr: str, output_dir: str) -> str:
    """Generate a single v3 Outlook-ready Word document."""

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)
    # Set narrow margins (more like email)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    name = contact["name"]
    title = contact["title"]
    company = contact["company"]
    panel = sanitize_text(contact["panel"])
    emails = contact["emails"]
    gender = GENDER_LOOKUP.get(contact["filename"], "M")

    # ---- Document header (reference only, light gray) ----
    header_p = doc.add_paragraph()
    set_paragraph_spacing(header_p, before_pt=0, after_pt=2)
    add_run(header_p, name, size=14, bold=True, color="333333")

    info_parts = []
    if title:
        info_parts.append(title)
    if company:
        info_parts.append(company)
    if email_addr:
        info_parts.append(email_addr)
    info_p = doc.add_paragraph()
    set_paragraph_spacing(info_p, before_pt=0, after_pt=2)
    add_run(info_p, " | ".join(info_parts), size=9, color="888888")

    panel_p = doc.add_paragraph()
    set_paragraph_spacing(panel_p, before_pt=0, after_pt=0)
    add_run(panel_p, f"Panel: {panel}", size=9, color="AAAAAA")

    # ---- Three email blocks ----
    for i, email_data in enumerate(emails):
        add_horizontal_line(doc)
        add_email_block(
            doc,
            email_num=i,
            subject=email_data["subject"],
            body=email_data["body"],
            recipient_email=email_addr,
            gender=gender,
            email_idx=i,
        )

    filename = f"{contact['filename']}_v3.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def main():
    outputs_dir = os.path.join(_ADHOC_ROOT, "outputs")
    target_run = "2026-04-16_12-46-26_adhoc_linkedin_retail_summit_2026"
    run_dir = os.path.join(outputs_dir, target_run)
    docs_dir = os.path.join(run_dir, "DOCS")
    os.makedirs(docs_dir, exist_ok=True)

    email_lookup = load_email_lookup(run_dir)
    print(f"Loaded {len(email_lookup)} email mappings from Apollo data.")

    generated = []
    for contact in CONTACTS_V2:
        fn = contact["filename"]
        email_addr = email_lookup.get(fn, "")
        if not email_addr:
            name_key = contact["name"].lower().replace(" ", "_")
            nfkd = unicodedata.normalize("NFKD", name_key)
            ascii_key = "".join(c for c in nfkd if not unicodedata.combining(c))
            email_addr = email_lookup.get(ascii_key, email_lookup.get(name_key, ""))

        filepath = generate_v3_doc(contact, email_addr, docs_dir)
        print(f"  \u2713 {os.path.basename(filepath)}  ({email_addr or 'NO EMAIL'})")
        generated.append(filepath)

    print(f"\nDone. {len(generated)} v3 docs saved to:\n  {docs_dir}")


if __name__ == "__main__":
    main()
