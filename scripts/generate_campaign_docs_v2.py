#!/usr/bin/env python3
"""
Generate v2 Word documents for 2 real outbound campaigns.
Humanized + copy-tuned versions.
Damian Hucz (AB Bechcicki) and Monika Sitkowska (ALDI Polska).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing_pt=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_spacing_pt:
        spacing.set(qn("w:line"), str(int(line_spacing_pt * 20)))
        spacing.set(qn("w:lineRule"), "exact")


def add_run(paragraph, text, size=11, bold=False, color="000000", font="Aptos"):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=12, after_pt=12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=6)
    add_run(p, text, size=16, bold=True, color="1F2937")


def add_section_header(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=14, after_pt=4)
    add_run(p, text, size=13, bold=True, color="1E40AF")


def add_subsection_header(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=10, after_pt=3)
    add_run(p, text, size=11, bold=True, color="374151")


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=1, after_pt=1)
    add_run(p, f"{label}: ", size=11, bold=True, color="4B5563")
    add_run(p, value, size=11, color="111827")


def add_body_text(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=2, after_pt=2, line_spacing_pt=15)
    add_run(p, text, size=11, color="111827")


def add_email_block(doc, label, subject, body):
    add_subsection_header(doc, label)
    add_label_value(doc, "Temat", subject)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=4, after_pt=4, line_spacing_pt=15)
    lines = body.split("\n")
    first = True
    for line in lines:
        if first:
            add_run(p, line, size=11, color="111827")
            first = False
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before_pt=2, after_pt=2, line_spacing_pt=15)
            add_run(p, line, size=11, color="111827")


def add_quote_block(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=4, after_pt=4, line_spacing_pt=15)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    add_run(p, text, size=11, color="4B5563")


# ─────────────────────────────────────────────────────────────────────────────
# CAMPAIGN DATA — v2 (humanized)
# ─────────────────────────────────────────────────────────────────────────────

DAMIAN_HUCZ = {
    "name": "Damian Hucz",
    "title": "Purchasing Director",
    "company": "AB Bechcicki Sp. z o.o.",
    "industry": "Dystrybucja materiałów budowlanych (hurtownia B2B)",
    "icp_fit": "Wysoki - Purchasing Director w dużej hurtowni budowlanej z szerokim portfelem dostawców i wieloma kategoriami zakupowymi.",
    "tier": "Tier 2 - Procurement Management / Dyrektorzy zakupów",
    "tier_rationale": (
        "Damian Hucz jako Purchasing Director odpowiada za zakupy w AB Bechcicki - dużej hurtowni materiałów "
        "budowlanych współpracującej z szerokim gronem producentów. To klasyczny Tier 2: lider funkcji zakupowej "
        "odpowiedzialny za warunki handlowe, pracę zespołu i wynik zakupowy w skali firmy."
    ),
    "tensions": (
        "AB Bechcicki to hurtownia materiałów budowlanych z wieloma kategoriami i dużą liczbą dostawców - "
        "producenci chemii, izolacji, farb, suchej zabudowy, dachów. Kluczowe napięcia:\n"
        "- Producenci materiałów budowlanych regularnie uzasadniają podwyżki zmianami kosztów surowców i energii.\n"
        "- Przy wielu kategoriach i dostawcach trudno utrzymać jednolity sposób oceny, co jest uzasadnione, a co nie.\n"
        "- Marża dystrybutora jest z natury niska - każda zmiana warunków ma bezpośredni wpływ na wynik.\n"
        "- Sezonowość budownictwa sprawia, że negocjacje warunków mają swój naturalny rytm i deadliny.\n"
        "- Różni kupcy przygotowują się do rozmów z dostawcami na różnym poziomie."
    ),
    "angle": (
        "Wspólna logika przygotowania do rozmów z dostawcami przy dużej liczbie kategorii. "
        "Gdy firma pracuje z wieloma producentami jednocześnie, łatwo o sytuację, w której jakość przygotowania "
        "jest nierówna - a to kosztuje."
    ),
    "hypothesis": (
        "Domyślam się, że przy tak szerokim gronie dostawców materiałów budowlanych - od chemii, przez izolacje, "
        "po farby i suchą zabudowę - jednym z ważniejszych pytań może być to, jak zapewnić, żeby zespół zakupowy "
        "przygotowywał się do rozmów z dostawcami w sposób powtarzalny i spójny. Przy niskiej marży dystrybutora "
        "nierówna jakość przygotowania pomiędzy kategoriami bywa kosztowna."
    ),

    # Email 1
    "e1_subject": "Pytanie o rozmowy z dostawcami - AB Bechcicki",
    "e1_body": (
        "Dzień dobry Panie Damianie,\n\n"
        "patrząc na rynek materiałów budowlanych i dystrybucji - domyślam się, że przy tak szerokim gronie dostawców "
        "jednym z trudniejszych tematów jest dziś spójne przygotowanie zespołu do rozmów z producentami. "
        "Szczególnie gdy podwyżki uzasadniane surowcami czy energią dotyczą wielu kategorii naraz.\n\n"
        "Jestem Tomasz Uściński z Profitii. Pomagamy dyrektorom zakupów w firmach dystrybucyjnych wypracować "
        "wspólną logikę oceny ofert i podwyżek, tak żeby wynik rozmów z dostawcami mniej zależał od tego, "
        "kto akurat prowadzi daną kategorię.\n\n"
        "Jeśli to dla Pana aktualny temat, proszę wybrać termin tutaj: [link do Calendly] - "
        "albo odpisać \"TAK\" i podać numer telefonu."
    ),

    # Follow-up 1
    "fu1_subject": "RE: Pytanie o rozmowy z dostawcami - AB Bechcicki",
    "fu1_body": (
        "Dzień dobry Panie Damianie,\n\n"
        "w dystrybucji budowlanej najtrudniejszy moment pojawia się, gdy kilku dostawców jednocześnie przychodzi "
        "z podwyżkami - i trzeba szybko ocenić, które z nich wynikają z realnych zmian kosztowych, a które "
        "warto zakwestionować.\n\n"
        "Pomagamy zespołom zakupowym zbudować prostą logikę weryfikacji takich argumentów - "
        "żeby rozmowa z dostawcą opierała się na danych, a nie tylko na przeczuciu.\n\n"
        "Jeśli chce Pan, mogę pokazać, jak to wygląda w praktyce - "
        "tutaj termin: [link do Calendly], albo odpisać \"TAK\" i podać numer telefonu."
    ),

    # Follow-up 2
    "fu2_subject": "RE: Pytanie o rozmowy z dostawcami - AB Bechcicki",
    "fu2_body": (
        "Dzień dobry Panie Damianie,\n\n"
        "z naszych obserwacji wynika, że w dystrybucji często najtrudniejsze nie jest samo negocjowanie - "
        "ale pokazanie zarządowi, w których kategoriach udało się uniknąć podwyżki, a gdzie jest jeszcze przestrzeń "
        "do poprawy warunków. Często zaczynamy od jednej kategorii - jako pilotaż - żeby sprawdzić, "
        "czy takie podejście ma sens w konkretnej firmie.\n\n"
        "Jeśli chce Pan, możemy to omówić na przykładzie - "
        "termin: [link do Calendly], albo odpisać \"TAK\" i podać numer telefonu."
    ),

    "rationale": (
        "DLA OSOBY: Damian Hucz jako Purchasing Director odpowiada za wynik zakupowy firmy - "
        "a przy dużej liczbie dostawców i kategorii wyzwanie polega na tym, żeby jakość przygotowania "
        "do rozmów była spójna, niezależnie od kupca czy kategorii.\n\n"
        "DLA FIRMY: AB Bechcicki to dystrybutor z niską marżą - każda zmiana warunków zakupowych "
        "przekłada się bezpośrednio na wynik. Przy szerokim portfelu (chemia, izolacje, farby, dachy) "
        "wspólna logika pracy z dostawcami ma wymierny efekt.\n\n"
        "DLA BRANŻY: Dystrybucja materiałów budowlanych mierzy się z podwyżkami producentów, "
        "sezonowością i konsolidacją rynku. Pytanie, ile z uzasadnień podwyżkowych "
        "odzwierciedla realia, a ile jest narracja.\n\n"
        "CO ODRÓŻNIA: Kampania mówi o realiach hurtowni budowlanej, nie o \"zakupach\" ogólnie. "
        "Każdy mail ma inny kąt: wspólna logika przygotowania → weryfikacja podwyżek → pilotaż i raportowanie."
    ),
}

MONIKA_SITKOWSKA = {
    "name": "Monika Sitkowska",
    "title": "Senior Director Purchasing / Dyrektor ds. Zakupów",
    "company": "ALDI Polska",
    "industry": "Retail / Sieć dyskontowa (FMCG, private label)",
    "icp_fit": "Wysoki - Senior Director Purchasing w sieci dyskontowej opartej na modelu EDLP, z silnym naciskiem na private label i cenę na półce.",
    "tier": "Tier 2 - Procurement Management / Dyrektorzy zakupów",
    "tier_rationale": (
        "Monika Sitkowska jako Senior Director Purchasing odpowiada za zakupy w ALDI Polska - "
        "sieci dyskontowej, w której efektywność zakupowa jest fundamentem modelu biznesowego. "
        "To Tier 2: lider funkcji zakupowej odpowiedzialny za warunki handlowe, negocjacje z dostawcami "
        "i wynik zakupowy. Seniority wysoka, ale rola jest procurement-management, nie C-Level."
    ),
    "tensions": (
        "ALDI to dyskont z modelem EDLP - niska cena na półce każdego dnia to DNA firmy. "
        "Kluczowe napięcia:\n"
        "- W dyskoncie cena na półce jest fundamentem przewagi - każda zaakceptowana podwyżka, "
        "która nie jest uzasadniona, osłabia pozycję wobec Lidla czy Biedronki.\n"
        "- Private label to rdzeń strategii ALDI - a tam negocjacje z dostawcami są intensywne, "
        "bo specyfikacja i cena są po stronie sieci.\n"
        "- Dostawcy FMCG regularnie uzasadniają podwyżki kosztami surowców, opakowań i logistyki.\n"
        "- Presja na cenę na półce wymaga, żeby ocena podwyżek była szybka, spójna i oparta na faktach.\n"
        "- Zarząd oczekuje mierzalnych efektów - ograniczonych podwyżek, obronionych warunków."
    ),
    "angle": (
        "Weryfikacja podwyżek dostawców FMCG i private label w modelu dyskontowym. "
        "Gdy cena na półce jest Twoim fundamentem, każda niezasadna podwyżka to strata przewagi."
    ),
    "hypothesis": (
        "Domyślam się, że w modelu dyskontowym EDLP, gdzie cena na półce jest fundamentem przewagi, "
        "szczególnie istotne jest dziś szybkie oddzielenie podwyżek dostawców FMCG, które wynikają "
        "z realnych zmian kosztowych, od tych, z którymi warto jeszcze porozmawiać. "
        "Przy skali operacji ALDI i presji ze strony konkurencji nierówna ocena podwyżek w portfelu "
        "kategorii może kosztować więcej, niż widać na poziomie pojedynczej rozmowy."
    ),

    # Email 1
    "e1_subject": "Pytanie o podwyżki dostawców FMCG - ALDI Polska",
    "e1_body": (
        "Dzień dobry Pani Moniko,\n\n"
        "w modelu dyskontowym - gdzie niska cena na półce jest fundamentem, a private label stanowi rdzeń oferty - "
        "domyślam się, że weryfikacja podwyżek od dostawców FMCG to dziś jeden z ważniejszych tematów. "
        "Szczególnie gdy uzasadnienia opierają się na surowcach, opakowaniach czy logistyce, "
        "a każda zaakceptowana podwyżka wpływa na pozycję cenową wobec Lidla czy Biedronki.\n\n"
        "Jestem Tomasz Uściński z Profitii. Pomagamy dyrektorom zakupów w sieciach handlowych "
        "szybciej oddzielać podwyżki realnie wynikające ze zmian kosztowych od tych, "
        "z którymi warto wrócić do rozmowy.\n\n"
        "Jeśli to dla Pani aktualny temat, proszę wybrać termin tutaj: [link do Calendly] - "
        "albo odpisać \"TAK\" i podać numer telefonu."
    ),

    # Follow-up 1
    "fu1_subject": "RE: Pytanie o podwyżki dostawców FMCG - ALDI Polska",
    "fu1_body": (
        "Dzień dobry Pani Moniko,\n\n"
        "w kontekście private label najtrudniejsze bywa nie samo odrzucenie podwyżki, ale zbudowanie "
        "spójnej logiki oceny - żeby cały zespół weryfikował argumenty dostawców w podobny sposób. "
        "Surowce, opakowania, transport - to realne koszty, ale nie każda podwyżka wynika z nich w takim stopniu, "
        "jak deklaruje dostawca.\n\n"
        "Pomagamy przełożyć dane o tym, co faktycznie dzieje się na rynku, na konkretną argumentację "
        "do rozmów z dostawcami.\n\n"
        "Jeśli chce Pani, mogę pokazać, jak to wygląda w praktyce - "
        "tutaj termin: [link do Calendly], albo odpisać \"TAK\" i podać numer telefonu."
    ),

    # Follow-up 2
    "fu2_subject": "RE: Pytanie o podwyżki dostawców FMCG - ALDI Polska",
    "fu2_body": (
        "Dzień dobry Pani Moniko,\n\n"
        "z naszych obserwacji wynika, że w sieciach handlowych kluczowy moment pojawia się, gdy zarząd pyta: "
        "\"ile podwyżek udało się ograniczyć i gdzie jest jeszcze przestrzeń?\" - "
        "a odpowiedź nie jest łatwa bez wspólnej logiki patrzenia na portfel kategorii.\n\n"
        "Często zaczynamy od jednej kategorii private label jako pilotażu - "
        "żeby sprawdzić, czy takie podejście pasuje do sposobu pracy w konkretnej sieci.\n\n"
        "Jeśli chce Pani, możemy to omówić na przykładzie - "
        "termin: [link do Calendly], albo odpisać \"TAK\" i podać numer telefonu."
    ),

    "rationale": (
        "DLA OSOBY: Monika Sitkowska jako Senior Director Purchasing odpowiada za efektywność zakupową "
        "ALDI Polska. Przy jej seniority kluczowe jest nie samo negocjowanie, ale zapewnienie "
        "spójnego podejścia do oceny podwyżek i pokazanie efektów zarządowi.\n\n"
        "DLA FIRMY: ALDI to dyskont EDLP - niska cena na półce każdego dnia. Każda niezasadna podwyżka "
        "osłabia pozycję cenową wobec Lidla i Biedronki. Private label jako rdzeń oferty "
        "wymaga szczególnie ostrych negocjacji z dostawcami.\n\n"
        "DLA BRANŻY: Retail FMCG w Polsce mierzy się z podwyżkami uzasadnianymi inflacją i kosztami. "
        "Część jest zasadna, część nie. Dyskonty z modelem EDLP są na to najbardziej wrażliwe.\n\n"
        "CO ODRÓŻNIA: Kampania od pierwszego zdania mówi językiem dyskontu - "
        "EDLP, private label, cena na półce, Lidl i Biedronka jako kontekst konkurencyjny. "
        "Każdy mail ma inny kąt: weryfikacja podwyżek → spójna logika oceny → pilotaż i raportowanie."
    ),
}


def build_campaign_doc(data, output_path):
    """Build a Word document for one campaign."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title(doc, f"Kampania outbound: {data['name']}")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=2)
    add_run(p, f"{data['title']} | {data['company']}", size=11, color="6B7280")
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before_pt=0, after_pt=2)
    add_run(p2, "Wersja: v2 (humanized + copy-tuned)", size=10, color="9CA3AF")

    add_horizontal_line(doc)

    add_section_header(doc, "1. Dane kampanii")
    add_label_value(doc, "Imię i nazwisko", data["name"])
    add_label_value(doc, "Stanowisko", data["title"])
    add_label_value(doc, "Firma", data["company"])
    add_label_value(doc, "Branża", data["industry"])
    add_label_value(doc, "Nadawca", "Tomasz Uściński, Head of Sales, Profitia")

    add_section_header(doc, "2. ICP i Tier")
    add_label_value(doc, "ICP fit", data["icp_fit"])
    add_label_value(doc, "Tier", data["tier"])
    add_subsection_header(doc, "Uzasadnienie Tieru")
    add_body_text(doc, data["tier_rationale"])

    add_section_header(doc, "3. Kontekst firmy i branży")
    add_subsection_header(doc, "Główne napięcia biznesowe")
    add_body_text(doc, data["tensions"])
    add_subsection_header(doc, "Messaging angle")
    add_body_text(doc, data["angle"])

    add_section_header(doc, "4. Hipoteza biznesowa")
    add_quote_block(doc, data["hypothesis"])

    add_section_header(doc, "5. Email 1")
    add_email_block(doc, "EMAIL 1", data["e1_subject"], data["e1_body"])

    add_section_header(doc, "6. Follow-up 1 (step 2)")
    add_email_block(doc, "FOLLOW-UP 1", data["fu1_subject"], data["fu1_body"])

    add_section_header(doc, "7. Follow-up 2 (step 3)")
    add_email_block(doc, "FOLLOW-UP 2", data["fu2_subject"], data["fu2_body"])

    add_section_header(doc, "8. Uzasadnienie kampanii")
    add_body_text(doc, data["rationale"])

    doc.save(output_path)
    print(f"  Zapisano: {output_path}")


def build_comparison_doc(output_path):
    """Build comparison summary document - v2."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title(doc, "Porównanie kampanii outbound - v2 (humanized)")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=2)
    add_run(p, "Damian Hucz (AB Bechcicki) vs Monika Sitkowska (ALDI Polska)", size=11, color="6B7280")

    add_horizontal_line(doc)

    add_section_header(doc, "1. Porównanie profili")

    for label, val_d, val_m in [
        ("Stanowisko", "Purchasing Director", "Senior Director Purchasing"),
        ("Firma", "AB Bechcicki Sp. z o.o.", "ALDI Polska"),
        ("Branża", "Dystrybucja materiałów budowlanych (B2B)", "Retail dyskontowy (FMCG, private label)"),
        ("Tier", "Tier 2 - Procurement Management", "Tier 2 - Procurement Management"),
        ("Model biznesowy", "Hurtownia B2B, niska marża dystrybucyjna", "Dyskont EDLP, cena na półce jako fundament"),
        ("Główny pain point", "Spójna logika przygotowania do rozmów przy wielu dostawcach", "Weryfikacja podwyżek dostawców FMCG w modelu dyskontowym"),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=2, after_pt=2)
        add_run(p, f"{label}:", size=11, bold=True, color="4B5563")
        p2 = doc.add_paragraph()
        set_paragraph_spacing(p2, before_pt=0, after_pt=0)
        add_run(p2, f"  Hucz: ", size=10, bold=True, color="1E40AF")
        add_run(p2, val_d, size=10, color="374151")
        p3 = doc.add_paragraph()
        set_paragraph_spacing(p3, before_pt=0, after_pt=2)
        add_run(p3, f"  Sitkowska: ", size=10, bold=True, color="1E40AF")
        add_run(p3, val_m, size=10, color="374151")

    add_horizontal_line(doc)

    add_section_header(doc, "2. Kluczowe różnice między kampaniami")

    differences = [
        (
            "Kontekst branżowy",
            "Hucz: Dystrybucja materiałów budowlanych B2B - kupuje od producentów, sprzedaje firmom budowlanym. "
            "Marża niska, sezonowość wysoka. Sitkowska: Dyskont FMCG z modelem EDLP - kupuje od producentów "
            "żywności i chemii, sprzedaje konsumentom. Cena na półce = fundament modelu."
        ),
        (
            "Messaging angle",
            "Hucz: Wspólna logika przygotowania do rozmów z dostawcami przy wielu kategoriach. "
            "Sitkowska: Weryfikacja podwyżek dostawców FMCG i private label - bo w dyskoncie każda "
            "niezasadna podwyżka to utrata przewagi cenowej."
        ),
        (
            "Sekwencja follow-upów",
            "Hucz: Email 1 = wspólna logika → FU1 = weryfikacja podwyżek → FU2 = pilotaż + raportowanie do zarządu. "
            "Sitkowska: Email 1 = weryfikacja podwyżek FMCG → FU1 = spójna logika oceny private label → FU2 = pilotaż + raportowanie."
        ),
        (
            "Specyfika językowa",
            "Hucz: Producenci materiałów budowlanych, kategorie budowlane, sezonowość. "
            "Sitkowska: Model EDLP, private label, cena na półce, Lidl/Biedronka jako kontekst konkurencyjny."
        ),
    ]

    for title, desc in differences:
        add_subsection_header(doc, title)
        add_body_text(doc, desc)

    add_horizontal_line(doc)

    add_section_header(doc, "3. Co się zmieniło vs v1")

    add_subsection_header(doc, "Damian Hucz - zmiany w v2")
    add_body_text(
        doc,
        "- Złagodzony ton pewności: zamiast \"300+ dostawców\" i konkretnych twierdzeń o portfelu, "
        "teraz \"szerokie grono dostawców\", \"wiele kategorii\", \"domyślam się\". "
        "- Mniej terminów technokratycznych: \"savings delivery\" zastąpione \"wynikiem rozmów z dostawcami\", "
        "\"cost drivers\" zastąpione \"co napędza koszty\". "
        "- Email 1 brzmi bardziej jak obserwacja z rynku, mniej jak analiza wewnętrzna. "
        "- CTA krótsze, bardziej naturalne."
    )

    add_subsection_header(doc, "Monika Sitkowska - zmiany w v2")
    add_body_text(
        doc,
        "- Kontekst ALDI od pierwszego zdania: model dyskontowy, EDLP, private label, "
        "cena na półce, Lidl i Biedronka pojawiają się w Email 1. "
        "- Mniej technokratycznego języka: \"avoided cost\" zastąpione prostszym sformułowaniem "
        "\"ile podwyżek udało się ograniczyć\". "
        "- FU1 mówi wprost o private label (zamiast ogólnego \"FMCG\"). "
        "- Ton bardziej ludzki: \"domyślam się\" zamiast twierdzeń, \"z naszych obserwacji\" "
        "zamiast eksperckiego pewnika."
    )

    add_horizontal_line(doc)

    add_section_header(doc, "4. Podsumowanie")
    add_body_text(
        doc,
        "Obie kampanie zachowują trafność merytoryczną z v1, ale brzmią bardziej naturalnie. "
        "Język jest lżejszy, mniej \"frameworkowy\", mniej technokratyczny. "
        "Każda kampania jest osadzona w realiach konkretnej firmy i branży, a różnice "
        "między Damianem a Moniką wynikają z kontekstu biznesowego, nie z szablonu."
    )

    doc.save(output_path)
    print(f"  Zapisano: {output_path}")


def main():
    output_dir = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..", "outputs", "word_campaigns"
    )
    os.makedirs(output_dir, exist_ok=True)

    print("Generuję kampanie outbound v2 (humanized)...\n")

    path1 = os.path.join(output_dir, "Damian_Hucz_campaign_v2.docx")
    build_campaign_doc(DAMIAN_HUCZ, path1)

    path2 = os.path.join(output_dir, "Monika_Sitkowska_campaign_v2.docx")
    build_campaign_doc(MONIKA_SITKOWSKA, path2)

    path3 = os.path.join(output_dir, "campaign_summary_comparison_v2.docx")
    build_comparison_doc(path3)

    print(f"\nGotowe. Pliki w: {output_dir}")


if __name__ == "__main__":
    main()
