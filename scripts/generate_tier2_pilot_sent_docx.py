#!/usr/bin/env python3
"""Generate final Word document for Tier 2 pilot sent version."""
import json
import os
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Load data
with open(os.path.join(ROOT, "tests", "output", "tier2_real_pilot_campaigns.json"), "r", encoding="utf-8") as f:
    campaigns = json.load(f)

with open(os.path.join(ROOT, "tests", "output", "tier2_real_pilot_apollo_send_result.json"), "r", encoding="utf-8") as f:
    send_result = json.load(f)

doc = Document()

# Style
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Title
title = doc.add_heading("Tier 2 Real Pilot - Sent Version", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Meta
meta = doc.add_paragraph()
meta.add_run("Data wysylki: ").bold = True
meta.add_run(f"{datetime.now().strftime('%Y-%m-%d')}\n")
meta.add_run("Mailbox: ").bold = True
meta.add_run(f"{send_result['mailbox_email']}\n")
meta.add_run("Cadence: ").bold = True
meta.add_run("D0 / D+2 / D+2\n")
meta.add_run("Calendly: ").bold = True
meta.add_run(f"{campaigns['calendly_url']}\n")
meta.add_run("Wersja copy: ").bold = True
meta.add_run(f"{campaigns['version']}\n")
meta.add_run("Overall verdict: ").bold = True
meta.add_run(f"{send_result['overall_verdict']}")

doc.add_paragraph("")

# Per-contact sections
for camp_key, camp in campaigns["campaigns"].items():
    contact = camp["contact"]
    sr = send_result["contacts"].get(camp_key, {})

    doc.add_heading(f"{contact['first_name']} {contact['last_name']} / {contact['company']}", level=2)

    # Contact info
    info = doc.add_paragraph()
    info.add_run("Email: ").bold = True
    info.add_run(f"{contact['email']}\n")
    info.add_run("Stanowisko: ").bold = True
    info.add_run(f"{contact['title']}\n")
    info.add_run("Tier: ").bold = True
    info.add_run(f"{camp['campaign_metadata']['tier']}\n")
    info.add_run("Angle: ").bold = True
    info.add_run(f"{camp['campaign_metadata']['angle']}\n")
    info.add_run("Sekwencja Apollo: ").bold = True
    info.add_run(f"{sr.get('sequence_name', 'n/a')}\n")
    info.add_run("Enrollment: ").bold = True
    info.add_run(f"{sr.get('enrollment_status', 'n/a')}\n")
    info.add_run("Verdict: ").bold = True
    verdict_run = info.add_run(f"{sr.get('verdict', 'n/a')}")
    if "BLOCKED" in sr.get("verdict", ""):
        verdict_run.font.color.rgb = RGBColor(200, 0, 0)

    # Emails
    for step_key, step_label in [("email_1", "EMAIL 1 (D0)"), ("follow_up_1", "FOLLOW-UP 1 (D+2)"), ("follow_up_2", "FOLLOW-UP 2 (D+4)")]:
        step = camp[step_key]

        doc.add_heading(step_label, level=3)

        subj_p = doc.add_paragraph()
        subj_p.add_run("Temat: ").bold = True
        subj_p.add_run(step["subject"])

        body_p = doc.add_paragraph()
        body_p.add_run("Tresc:\n").bold = True
        body_run = body_p.add_run(step["body"])
        body_run.font.size = Pt(10)
        body_run.font.color.rgb = RGBColor(50, 50, 50)

    doc.add_page_break()

# Save
output_path = os.path.join(ROOT, "outputs", "word_campaigns", "tier2_real_pilot_sent_version.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Saved: {output_path}")
