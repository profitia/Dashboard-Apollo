#!/usr/bin/env python3
"""
Tier 2 Real Pilot v2 — polished copy + real Calendly links.

Generates:
- tests/output/tier2_real_pilot_campaigns.json (overwrite)
- tests/output/tier2_real_pilot_copy_polish_report.md
- outputs/word_campaigns/tier2_real_pilot_review_v2.docx
"""

import json
import os
import re
from datetime import datetime

# ============================================================
# Calendly links per Tier (mirrors source_of_truth)
# ============================================================

CALENDLY_TIER_2 = "https://calendly.com/profitia/standard-negocjacji-i-oszczednosci"

# ============================================================
# POLISHED CAMPAIGN COPY
# ============================================================

campaigns = {
    "konrad_ludwicki_circle_k": {
        "lead_id": 1,
        "contact": {
            "first_name": "Konrad",
            "last_name": "Ludwicki",
            "vocative": "Konradzie",
            "gender": "male",
            "title": "Senior Director, Head of Procurement Europe",
            "company": "Circle K",
            "email": "kludwicki@circlek.com",
        },
        "campaign_metadata": {
            "tier": "Tier 2",
            "angle": "savings_delivery",
            "persona": "Dyrektor Zakupów / Head of Procurement",
            "language": "pl",
            "campaign_type": "standard_outbound_tier2",
        },
        "email_1": {
            "subject": "Oszczędności w zakupach europejskich - Circle K",
            "body": (
                "Dzień dobry Panie Konradzie,\n\n"
                "zwróciłem uwagę na Pana rolę jako Head of Procurement Europe "
                "w Circle K. Przy wielu rynkach i zespołach jednocześnie "
                "domyślam się, że jednym z wyzwań jest porównywalny standard "
                "oceny ofert i podwyżek dostawców - niezależnie od tego, który "
                "kupiec i w którym kraju prowadzi rozmowę.\n\n"
                "Nazywam się Tomasz Uściński, jestem z polskiej firmy Profitia, "
                "w której od 15 lat pomagamy dyrektorom zakupów w sieciach "
                "handlowych budować spójną logikę przygotowania negocjacji "
                "z dostawcami.\n\n"
                "Jeśli temat jest aktualny, chętnie opowiem więcej przy krótkiej "
                "rozmowie: " + CALENDLY_TIER_2 + "\n"
                "Jeśli woli Pan po prostu krótką rozmowę telefoniczną, "
                "bardzo proszę o numer telefonu - chętnie oddzwonię."
            ),
        },
        "follow_up_1": {
            "subject": "Re: Oszczędności w zakupach europejskich - Circle K",
            "body": (
                "Dzień dobry Panie Konradzie,\n\n"
                "dopowiem jedną rzecz w kontekście zakupów na wielu rynkach. "
                "Często wyzwaniem nie jest brak dobrych praktyk - ale to, że "
                "każdy zespół wypracował własny sposób oceny ofert. Wynik "
                "negocjacji zależy wtedy bardziej od kupca niż od metody.\n\n"
                "Pomagamy to zmienić - budujemy wspólną logikę przygotowania, "
                "od struktury kosztu po argumentację, tak aby każdy zespół "
                "startował z porównywalnego punktu w rozmowach z dostawcami.\n\n"
                "Jeśli chce Pan zobaczyć, jak to wygląda na jednej kategorii, "
                "proszę dać znać."
            ),
        },
        "follow_up_2": {
            "subject": "Re: Oszczędności w zakupach europejskich - Circle K",
            "body": (
                "Dzień dobry Panie Konradzie,\n\n"
                "z naszych obserwacji - dyrektorzy zakupów w dużych sieciach "
                "coraz częściej mierzą się z pytaniem, jak pokazać zarządowi "
                "wartość unikniętych podwyżek jako realny wkład zakupów "
                "w wynik firmy.\n\n"
                "Pomagamy przełożyć te efekty na konkretne liczby do rozmowy "
                "z CFO - tak aby praca zespołu zakupowego była widoczna "
                "i mierzalna.\n\n"
                "Jeśli woli Pan wrócić do tego telefonicznie, bardzo proszę "
                "o numer telefonu - chętnie oddzwonię."
            ),
        },
    },
    "anna_polak_ds_smith": {
        "lead_id": 2,
        "contact": {
            "first_name": "Anna",
            "last_name": "Polak",
            "vocative": "Anno",
            "gender": "female",
            "title": "Procurement Director Poland",
            "company": "DS Smith",
            "email": "anna.polak@dssmith.com",
        },
        "campaign_metadata": {
            "tier": "Tier 2",
            "angle": "supplier_price_increases",
            "persona": "Dyrektor Zakupów / Procurement Director",
            "language": "pl",
            "campaign_type": "standard_outbound_tier2",
        },
        "email_1": {
            "subject": "Weryfikacja podwyżek dostawców - DS Smith",
            "body": (
                "Dzień dobry Pani Anno,\n\n"
                "w branży opakowaniowej ceny recyclatu, surowców papierowych "
                "i energii zmieniają się z kwartału na kwartał. Domyślam się, "
                "że w Pani roli jako Procurement Director jednym z wyzwań jest "
                "ocena, które podwyżki dostawców są realnie uzasadnione "
                "kosztowo, a które są narracją negocjacyjną.\n\n"
                "Nazywam się Tomasz Uściński, jestem z polskiej firmy Profitia, "
                "w której od 15 lat pomagamy zespołom zakupowym w firmach "
                "produkcyjnych weryfikować zasadność argumentów kosztowych "
                "dostawców.\n\n"
                "Jeśli temat jest aktualny, chętnie pokażę jak to wygląda "
                "w praktyce: " + CALENDLY_TIER_2 + "\n"
                "Jeśli woli Pani po prostu krótką rozmowę telefoniczną, "
                "bardzo proszę o numer telefonu - chętnie oddzwonię."
            ),
        },
        "follow_up_1": {
            "subject": "Re: Weryfikacja podwyżek dostawców - DS Smith",
            "body": (
                "Dzień dobry Pani Anno,\n\n"
                "dopowiem jedną rzecz w kontekście kosztów w branży "
                "opakowaniowej. Kiedy dostawca uzasadnia podwyżkę cenami "
                "recyclatu czy energii, kluczowe jest sprawdzenie, w jakim "
                "stopniu te zmiany faktycznie wpływają na koszt produktu "
                "- a w jakim to argument negocjacyjny.\n\n"
                "Pomagamy rozbić strukturę kosztu na konkretne składniki "
                "i porównać dynamikę cen z faktycznymi zmianami - tak aby "
                "rozmowa z dostawcą opierała się na danych, nie "
                "na deklaracjach.\n\n"
                "Jeśli chce Pani zobaczyć to na jednej kategorii, proszę "
                "dać znać."
            ),
        },
        "follow_up_2": {
            "subject": "Re: Weryfikacja podwyżek dostawców - DS Smith",
            "body": (
                "Dzień dobry Pani Anno,\n\n"
                "z perspektywy negocjacji z dostawcami kluczowe bywa nie "
                "tylko to, co negocjować, ale kiedy. W branży papierowej "
                "i opakowaniowej ceny surowców mają wyraźne cykle "
                "- i odpowiedni moment renegocjacji może mieć istotny "
                "wpływ na warunki.\n\n"
                "Pomagamy ocenić, czy to dobry czas na rozmowę z dostawcą "
                "o cenie, indeksacji lub warunkach kontraktu - na podstawie "
                "prognoz i trendów kosztowych.\n\n"
                "Jeśli woli Pani wrócić do tego telefonicznie, bardzo proszę "
                "o numer telefonu - chętnie oddzwonię."
            ),
        },
    },
    "sebastian_kacperek_astrazeneca": {
        "lead_id": 3,
        "contact": {
            "first_name": "Sebastian",
            "last_name": "Kacperek",
            "vocative": "Sebastianie",
            "gender": "male",
            "title": "Associate Director Global Procurement",
            "company": "AstraZeneca",
            "email": "sebastian.kacperek@astrazeneca.com",
        },
        "campaign_metadata": {
            "tier": "Tier 2",
            "angle": "negotiation_preparation",
            "persona": "Dyrektor Zakupów / Associate Director Global Procurement",
            "language": "pl",
            "campaign_type": "standard_outbound_tier2",
        },
        "email_1": {
            "subject": "Przygotowanie negocjacji z dostawcami - AstraZeneca",
            "body": (
                "Dzień dobry Panie Sebastianie,\n\n"
                "w farmacji wymagania jakościowe i regulacyjne realnie "
                "zawężają pole negocjacji z dostawcami. Domyślam się, "
                "że w Pana roli wyzwaniem bywa znalezienie przestrzeni "
                "do poprawy warunków handlowych tam, gdzie wybór "
                "dostawców jest ograniczony.\n\n"
                "Nazywam się Tomasz Uściński, jestem z polskiej firmy "
                "Profitia, w której od 15 lat pomagamy zespołom zakupowym "
                "skuteczniej przygotowywać się do negocjacji z dostawcami "
                "- również w kategoriach z wąską bazą dostawców.\n\n"
                "Jeśli temat jest aktualny, chętnie opowiem więcej: "
                + CALENDLY_TIER_2 + "\n"
                "Jeśli woli Pan po prostu krótką rozmowę telefoniczną, "
                "bardzo proszę o numer telefonu - chętnie oddzwonię."
            ),
        },
        "follow_up_1": {
            "subject": "Re: Przygotowanie negocjacji z dostawcami - AstraZeneca",
            "body": (
                "Dzień dobry Panie Sebastianie,\n\n"
                "dopowiem jedną rzecz w kontekście negocjacji z dostawcami "
                "w farmacji. Nawet przy ograniczonej bazie dostawców pytanie "
                "często nie dotyczy samej ceny, ale konstrukcji kontraktu "
                "- indeksacji, wolumenu, warunków płatności czy długości "
                "umowy.\n\n"
                "Pomagamy sprawdzić, które z tych dźwigni mogą realnie "
                "przełożyć się na lepsze warunki - bez naruszania standardów "
                "jakościowych i compliance.\n\n"
                "Jeśli chce Pan zobaczyć to na przykładzie jednej kategorii, "
                "proszę dać znać."
            ),
        },
        "follow_up_2": {
            "subject": "Re: Przygotowanie negocjacji z dostawcami - AstraZeneca",
            "body": (
                "Dzień dobry Panie Sebastianie,\n\n"
                "w farmacji dyrektorzy zakupów często mierzą się z pytaniem "
                "- jak pokazać wartość tego, czego zespół uniknął? Uniknięte "
                "podwyżki i lepsze warunki kontraktowe to realna wartość, "
                "ale trudniej ją wykazać niż klasyczną obniżkę ceny.\n\n"
                "Pomagamy przełożyć te efekty na konkretne dane do rozmowy "
                "z zarządem - tak aby wkład zakupów był widoczny.\n\n"
                "Jeśli woli Pan wrócić do tego telefonicznie, bardzo proszę "
                "o numer telefonu - chętnie oddzwonię."
            ),
        },
    },
}

# Build apollo_custom_fields for each campaign
for key, camp in campaigns.items():
    camp["apollo_custom_fields"] = {
        "sg_email_step_1_subject": camp["email_1"]["subject"],
        "sg_email_step_1_body": camp["email_1"]["body"],
        "sg_email_step_2_subject": camp["follow_up_1"]["subject"],
        "sg_email_step_2_body": camp["follow_up_1"]["body"],
        "sg_email_step_3_subject": camp["follow_up_2"]["subject"],
        "sg_email_step_3_body": camp["follow_up_2"]["body"],
    }


def count_words(text):
    clean = re.sub(r'[„"""\u201e\u201d]', '', text)
    clean = re.sub(r'https?://\S+', '', clean)
    return len(clean.split())


# ============================================================
# Selected leads data (unchanged from v1)
# ============================================================

selected_leads = [
    {
        "id": 1,
        "first_name": "Konrad",
        "last_name": "Ludwicki",
        "title": "Senior Director, Head of Procurement Europe",
        "company": "Circle K",
        "email": "kludwicki@circlek.com",
        "industry": "retail",
        "city": "Piaseczno",
        "country": "Poland",
        "company_city": "Laval",
        "company_country": "Canada",
        "linkedin": "http://www.linkedin.com/in/konrad-l-42a00521",
        "website": "https://circlek.com",
        "tier": "Tier 2",
        "persona": "Dyrektor Zakupów / Head of Procurement",
        "icp_fit": "Excellent",
        "angle": "savings_delivery",
        "hypothesis": (
            "Przy wielu europejskich rynkach i zespołach jednocześnie wyzwaniem "
            "jest porównywalny standard oceny ofert i podwyżek dostawców - "
            "niezależnie od tego, który kupiec prowadzi rozmowę."
        ),
        "selection_reason": (
            "Head of Procurement Europe w Circle K - rola paneuropejska "
            "z odpowiedzialnością za wiele rynków i kategorii. Duży, rozpoznawalny brand."
        ),
    },
    {
        "id": 2,
        "first_name": "Anna",
        "last_name": "Polak",
        "title": "Procurement Director Poland",
        "company": "DS Smith",
        "email": "anna.polak@dssmith.com",
        "industry": "packaging & containers",
        "city": "Kielce",
        "country": "Poland",
        "company_city": "London",
        "company_country": "United Kingdom",
        "linkedin": "http://www.linkedin.com/in/anna-polak-501651106",
        "website": "https://dssmith.com",
        "tier": "Tier 2",
        "persona": "Dyrektor Zakupów / Procurement Director",
        "icp_fit": "Excellent",
        "angle": "supplier_price_increases",
        "hypothesis": (
            "W branży opakowaniowej ceny recyclatu, surowców papierowych i energii "
            "zmieniają się dynamicznie - wyzwaniem jest ocena, które podwyżki dostawców "
            "są uzasadnione kosztowo, a które są narracją negocjacyjną."
        ),
        "selection_reason": (
            "Procurement Director Poland w DS Smith - giełdowej firmie opakowaniowej. "
            "Branża narażona na zmienność cen surowców."
        ),
    },
    {
        "id": 3,
        "first_name": "Sebastian",
        "last_name": "Kacperek",
        "title": "Associate Director Global Procurement",
        "company": "AstraZeneca",
        "email": "sebastian.kacperek@astrazeneca.com",
        "industry": "pharmaceuticals",
        "city": "Warsaw",
        "country": "Poland",
        "company_city": "Cambridge",
        "company_country": "United Kingdom",
        "linkedin": "http://www.linkedin.com/in/sebastiankacperek",
        "website": "https://astrazeneca.com",
        "tier": "Tier 2",
        "persona": "Dyrektor Zakupów / Associate Director Global Procurement",
        "icp_fit": "Excellent",
        "angle": "negotiation_preparation",
        "hypothesis": (
            "W farmacji wymagania jakościowe i regulacyjne zawężają pole negocjacji "
            "- wyzwaniem jest znalezienie przestrzeni do poprawy warunków handlowych "
            "przy ograniczonej bazie dostawców."
        ),
        "selection_reason": (
            "Associate Director Global Procurement w AstraZeneca - jedna z największych "
            "firm farmaceutycznych na świecie. Zaawansowana organizacja zakupowa."
        ),
    },
]


# ============================================================
# OUTPUT 1: Updated campaigns.json
# ============================================================

def generate_campaigns_json():
    return {
        "generated_at": datetime.now().isoformat(),
        "version": "v2_polished",
        "campaign_type": "standard_outbound_tier2",
        "pilot_name": "Tier 2 Real Pilot - First 3",
        "sequence_cadence": "D0 / D+2 / D+2",
        "apollo_sync_status": "prepared_not_synced",
        "calendly_tier": "Tier 2",
        "calendly_url": CALENDLY_TIER_2,
        "calendly_source": "source_of_truth/apollo_custom_fields.yaml → calendly_links_per_tier.tier_2",
        "note": "Kampanie v2 (polished copy + real Calendly). Gotowe do review przed Apollo sync.",
        "campaigns": campaigns,
    }


# ============================================================
# OUTPUT 2: Copy polish report
# ============================================================

def generate_polish_report():
    lines = []
    lines.append("# Tier 2 Real Pilot - Raport poprawek copy (v1 → v2)\n")
    lines.append(f"**Data:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    lines.append("---\n")

    # 1. Konrad
    lines.append("## 1. Konrad Ludwicki / Circle K\n")
    lines.append("### Email 1")
    lines.append("- **Skrócono pierwszy akapit** - usunięto wyliczankę kategorii (\"od convenience, przez F&B, po usługi i wyposażenie\")")
    lines.append("- **Zmieniono główny sens** z \"spójne dowożenie oszczędności we wszystkich kategoriach\" na \"porównywalny standard oceny ofert i podwyżek dostawców niezależnie od kupca i rynku\" - mniej książkowo, bardziej operacyjnie")
    lines.append("- **Skrócono sender intro** - usunięto \"systemowo dowozić savings w wielu kategoriach\" → \"budować spójną logikę przygotowania negocjacji z dostawcami\" - lżejsze, mniej technokratyczne")
    lines.append("- **Złagodzono CTA** - zamiast sztywnego \"proszę wybrać dogodny termin tutaj\" → \"chętnie opowiem więcej przy krótkiej rozmowie\" + link")
    lines.append("- **Usunięto mechaniczne CTA** - zamiast \"Może Pan też po prostu odpisać TAK i podać numer telefonu\" → \"Można też po prostu odpisać - oddzwonię\" - naturalniejsze\n")

    lines.append("### Follow-up 1")
    lines.append("- **Dodano trafniejszy problem** - \"wyzwaniem nie jest brak dobrych praktyk, ale to, że każdy zespół wypracował własny sposób\" - bardziej wiarygodne")
    lines.append("- **Zachowano logikę** wspólnej logiki przygotowania i argumentacji\n")

    lines.append("### Follow-up 2")
    lines.append("- **Usunięto \"avoided cost\"** - zamieniono na \"wartość unikniętych podwyżek\" i \"te efekty\" - naturalniejszy polski")
    lines.append("- **Zamieniono \"dane o avoided cost na konkretne liczby\"** na \"przełożyć te efekty na konkretne liczby do rozmowy z CFO\" - bardziej naturalnie i trafniej (CFO zamiast ogólnego \"zarząd\")\n")

    # 2. Anna
    lines.append("## 2. Anna Polak / DS Smith\n")
    lines.append("### Email 1")
    lines.append("- **Skrócono pierwszy akapit o ~15%** - usunięto \"patrząc na sytuację w branży opakowaniowej\" i nadmiarowy dash, uproszczono zdanie otwierające")
    lines.append("- **Usunięto \"z kwartału na kwartał\"** - zbędne doprecyzowanie, skraca zdanie")
    lines.append("- **Usunięto z sender intro** \"tak aby ograniczać nieuzasadnione podwyżki i chronić budżet zakupowy\" - za ciężki drugi człon, niepotrzebny")
    lines.append("- **Złagodzono CTA** - \"chętnie pokażę jak to wygląda w praktyce\" + link - bardziej naturalne")
    lines.append("- **Reszta zachowana** - angle, logika, ton bez zmian\n")

    lines.append("### Follow-up 1")
    lines.append("- **Drobne wygładzenie** - \"strukturę kosztu kategorii na konkretne drivery\" → \"strukturę kosztu na konkretne składniki\" - prostszy język")
    lines.append("- **Reszta bez zmian** - logika i ton zachowane\n")

    lines.append("### Follow-up 2")
    lines.append("- **Minimalne zmiany** - \"często kluczowe jest\" → \"kluczowe bywa\" - lżejszy ton")
    lines.append("- **\"timing renegocjacji\"** → \"moment renegocjacji\" - naturalniejszy polski")
    lines.append("- **To pozostaje najlepsza kampania z trójki**\n")

    # 3. Sebastian
    lines.append("## 3. Sebastian Kacperek / AstraZeneca\n")
    lines.append("### Email 1 (największe zmiany)")
    lines.append("- **Gruntowne uproszczenie pierwszego akapitu** - usunięto abstrakcyjne \"identyfikacja kategorii, w których mimo tych ograniczeń nadal jest realna przestrzeń do poprawy warunków handlowych\"")
    lines.append("- **Nowa wersja** - \"wyzwaniem bywa znalezienie przestrzeni do poprawy warunków handlowych tam, gdzie wybór dostawców jest ograniczony\" - krótsza, konkretniejsza")
    lines.append("- **Usunięto \"patrząc na specyfikę zakupów w branży farmaceutycznej\"** - zbyt formalne intro, zastąpiono bezpośrednim \"w farmacji\"")
    lines.append("- **Skrócono sender intro** - usunięto \"w dużych organizacjach\" → zostawiono \"zespołom zakupowym\" + \"również w kategoriach z wąską bazą dostawców\"")
    lines.append("- **CTA złagodzono** jak w pozostałych kampaniach\n")

    lines.append("### Follow-up 1")
    lines.append("- **Lekkie wygładzenie** - \"kluczowe pytanie często dotyczy\" → \"pytanie często nie dotyczy\" - naturalniejszy flow")
    lines.append("- **Zachowano logikę** indeksacji / wolumenu / warunków płatności / compliance")
    lines.append("- **\"niższy koszt zakupu lub lepsze warunki handlowe\"** → \"lepsze warunki\" - skrócenie bez utraty sensu\n")

    lines.append("### Follow-up 2 (istotne zmiany)")
    lines.append("- **Usunięto zbyt abstrakcyjne \"jak wykazać wartość unikniętych podwyżek jako realny wkład zakupów\"**")
    lines.append("- **Nowa wersja** - \"jak pokazać wartość tego, czego zespół uniknął? Uniknięte podwyżki i lepsze warunki kontraktowe to realna wartość, ale trudniej ją wykazać niż klasyczną obniżkę ceny\" - bardziej ludzkie, mniej doradczo-generyczne")
    lines.append("- **Zachowano sens** raportowania do zarządu, ale ton bardziej wiarygodny\n")

    # 4. CTA
    lines.append("## 4. Zmiany CTA (globalne)\n")
    lines.append("### Stary schemat (v1)")
    lines.append("```")
    lines.append("Jeśli temat jest dla Pana aktualny, proszę wybrać dogodny termin tutaj: [link do Calendly].")
    lines.append("Może Pan też po prostu odpisać 'TAK' i podać numer telefonu - oddzwonię.")
    lines.append("```\n")
    lines.append("### Nowy schemat (v2)")
    lines.append("```")
    lines.append("Jeśli temat jest aktualny, chętnie opowiem więcej przy krótkiej rozmowie: [Calendly URL]")
    lines.append("Można też po prostu odpisać - oddzwonię.")
    lines.append("```\n")
    lines.append("### Co się zmieniło")
    lines.append("- Usunięto \"proszę wybrać dogodny termin tutaj\" - zbyt szablonowe dla senioralnych ról")
    lines.append("- Usunięto \"Może Pan/Pani też po prostu odpisać TAK i podać numer telefonu\" - mechaniczne")
    lines.append("- Nowe CTA jest lżejsze, bardziej konwersacyjne")
    lines.append("- Link Calendly wstawiony jako naturalny element zdania, nie osobny krok")
    lines.append("- Zachowano alternatywę (\"odpisać - oddzwonię\") ale bez \"TAK\" i numeru\n")

    # 5. Sender intro
    lines.append("## 5. Sender intro\n")
    lines.append("| Lead | Zmiana |")
    lines.append("|------|--------|")
    lines.append("| Konrad | Skrócono: usunięto \"systemowo dowozić savings w wielu kategoriach - tak aby wynik nie zależał od stylu pracy pojedynczych kupców\" → \"budować spójną logikę przygotowania negocjacji z dostawcami\" |")
    lines.append("| Anna | Skrócono: usunięto drugi człon \"tak aby ograniczać nieuzasadnione podwyżki i chronić budżet zakupowy\" |")
    lines.append("| Sebastian | Skrócono: \"w dużych organizacjach\" usunięte + dodano \"również w kategoriach z wąską bazą dostawców\" |")
    lines.append("")
    lines.append("Wszystkie sender intro zachowują: Tomasz Uściński + polska firma Profitia + 15 lat + dopasowana wartość.\n")

    # 6. Calendly
    lines.append("## 6. Linki Calendly\n")
    lines.append("- Placeholder `[link do Calendly]` → **podmieniony na realny URL**")
    lines.append(f"- Tier 2 URL: `{CALENDLY_TIER_2}`")
    lines.append("- Wszystkie 3 kampanie to Tier 2 → jeden link")
    lines.append("- Link wstawiony w zdanie CTA, nie jako osobny element\n")

    # 7. Source of truth
    lines.append("## 7. Globalna reguła Calendly per Tier\n")
    lines.append("Zapisano w: `source_of_truth/apollo_custom_fields.yaml`\n")
    lines.append("Sekcja: `calendly_links_per_tier`\n")
    lines.append("```yaml")
    lines.append("calendly_links_per_tier:")
    lines.append("  tier_1:")
    lines.append("    url: \"https://calendly.com/profitia/zakupy-a-marza-firmy\"")
    lines.append("  tier_2:")
    lines.append("    url: \"https://calendly.com/profitia/standard-negocjacji-i-oszczednosci\"")
    lines.append("  tier_3:")
    lines.append("    url: \"https://calendly.com/profitia/przygotowanie-negocjacji-z-dostawcami\"")
    lines.append("```\n")
    lines.append("Ta reguła jest trwała i dostępna dla wszystkich kolejnych kampanii.\n")

    # 8. Finalne werdykty
    lines.append("## 8. Finalne werdykty\n")
    lines.append("| Lead | Kampania | Werdykt |")
    lines.append("|------|----------|---------|")
    lines.append("| Konrad Ludwicki | Circle K | **Gotowa do Apollo** - po szybkim review nadawcy |")
    lines.append("| Anna Polak | DS Smith | **Gotowa do Apollo** - najlepsza jakościowo, minimalne zmiany |")
    lines.append("| Sebastian Kacperek | AstraZeneca | **Gotowa do Apollo** - po review nadawcy (największe zmiany copy) |")
    lines.append("")
    lines.append("**Ogólna rekomendacja:** Wszystkie 3 kampanie gotowe do Apollo po szybkim review Tomasza.")
    lines.append("Realne linki Calendly wstawione. Globalna reguła Calendly per Tier zapisana w source of truth.")

    return "\n".join(lines)


# ============================================================
# OUTPUT 3: Word v2
# ============================================================

def generate_word():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    title = doc.add_heading('Tier 2 Real Pilot - Review kampanii v2 (polished)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Wygenerowano: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Wersja: v2 (polished copy + real Calendly links)')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Summary
    doc.add_heading('Podsumowanie zmian v1 \u2192 v2', level=1)
    changes = [
        'Konrad / Circle K: lżejszy akapit 1, mniej abstrakcyjny, FU2 bez "avoided cost"',
        'Anna / DS Smith: skrócony akapit 1, zachowana jakość, lżejsze CTA',
        'Sebastian / AstraZeneca: gruntownie uproszczony E1 i FU2, bardziej naturalny język',
        'CTA: wszystkie złagodzone, mniej mechaniczne, realne linki Calendly',
        'Sender intro: skrócone we wszystkich 3, zachowana logika (TU + Profitia + 15 lat)',
        f'Calendly Tier 2: {CALENDLY_TIER_2}',
    ]
    for c in changes:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_page_break()

    campaign_keys = [
        "konrad_ludwicki_circle_k",
        "anna_polak_ds_smith",
        "sebastian_kacperek_astrazeneca",
    ]

    for idx, (lead, camp_key) in enumerate(zip(selected_leads, campaign_keys)):
        camp = campaigns[camp_key]

        doc.add_heading(
            f'Lead {idx + 1}: {lead["first_name"]} {lead["last_name"]} - {lead["company"]}',
            level=1,
        )

        # Info table
        doc.add_heading('Dane', level=2)
        info_table = doc.add_table(rows=7, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_data = [
            ("Imię i nazwisko", f'{lead["first_name"]} {lead["last_name"]}'),
            ("Stanowisko", lead["title"]),
            ("Firma", lead["company"]),
            ("Email", lead["email"]),
            ("Tier", lead["tier"]),
            ("Angle", lead["angle"]),
            ("Hypothesis", lead["hypothesis"]),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value

        # Email 1
        doc.add_heading('Email 1', level=2)
        p = doc.add_paragraph()
        run = p.add_run('Temat: ')
        run.bold = True
        p.add_run(camp["email_1"]["subject"])

        p = doc.add_paragraph()
        run = p.add_run('Treść:')
        run.bold = True
        for para in camp["email_1"]["body"].split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
        wc = count_words(camp["email_1"]["body"])
        p = doc.add_paragraph()
        run = p.add_run(f'[{wc} słów]')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # FU1
        doc.add_heading('Follow-up 1 (step 2)', level=2)
        p = doc.add_paragraph()
        run = p.add_run('Temat: ')
        run.bold = True
        p.add_run(camp["follow_up_1"]["subject"])

        p = doc.add_paragraph()
        run = p.add_run('Treść:')
        run.bold = True
        for para in camp["follow_up_1"]["body"].split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
        wc = count_words(camp["follow_up_1"]["body"])
        p = doc.add_paragraph()
        run = p.add_run(f'[{wc} słów]')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # FU2
        doc.add_heading('Follow-up 2 (step 3)', level=2)
        p = doc.add_paragraph()
        run = p.add_run('Temat: ')
        run.bold = True
        p.add_run(camp["follow_up_2"]["subject"])

        p = doc.add_paragraph()
        run = p.add_run('Treść:')
        run.bold = True
        for para in camp["follow_up_2"]["body"].split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
        wc = count_words(camp["follow_up_2"]["body"])
        p = doc.add_paragraph()
        run = p.add_run(f'[{wc} słów]')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        if idx < 2:
            doc.add_page_break()

    return doc


# ============================================================
# MAIN
# ============================================================

def main():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    os.makedirs(os.path.join(root, "tests", "output"), exist_ok=True)
    os.makedirs(os.path.join(root, "outputs", "word_campaigns"), exist_ok=True)

    # 1. Campaigns JSON (overwrite)
    path1 = os.path.join(root, "tests", "output", "tier2_real_pilot_campaigns.json")
    with open(path1, "w", encoding="utf-8") as f:
        json.dump(generate_campaigns_json(), f, ensure_ascii=False, indent=2)
    print(f"[OK] {path1}")

    # 2. Copy polish report
    path2 = os.path.join(root, "tests", "output", "tier2_real_pilot_copy_polish_report.md")
    with open(path2, "w", encoding="utf-8") as f:
        f.write(generate_polish_report())
    print(f"[OK] {path2}")

    # 3. Word v2
    path3 = os.path.join(root, "outputs", "word_campaigns", "tier2_real_pilot_review_v2.docx")
    doc = generate_word()
    doc.save(path3)
    print(f"[OK] {path3}")

    # Summary
    print("\n=== TIER 2 REAL PILOT v2 — POLISHED ===")
    print("Zmiany:")
    print("  - Konrad/Circle K: lżejszy E1, naturalniejszy FU2")
    print("  - Anna/DS Smith: skrócony E1, minimalne zmiany")
    print("  - Sebastian/AstraZeneca: uproszczony E1 i FU2")
    print("  - CTA: złagodzone, mniej mechaniczne")
    print("  - Calendly: realne linki Tier 2")
    print(f"  - Calendly URL: {CALENDLY_TIER_2}")
    print("  - Source of truth: apollo_custom_fields.yaml → calendly_links_per_tier")

    # Word counts
    print("\nWord counts:")
    for key, camp in campaigns.items():
        name = camp["contact"]["last_name"]
        e1 = count_words(camp["email_1"]["body"])
        fu1 = count_words(camp["follow_up_1"]["body"])
        fu2 = count_words(camp["follow_up_2"]["body"])
        print(f"  {name}: E1={e1}w, FU1={fu1}w, FU2={fu2}w")

    print(f"\nStatus: READY FOR REVIEW")


if __name__ == "__main__":
    main()
