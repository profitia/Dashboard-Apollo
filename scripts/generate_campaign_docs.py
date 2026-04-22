#!/usr/bin/env python3
"""
Generate Word documents for 2 real outbound campaigns.
Damian Hucz (AB Bechcicki) and Monika Sitkowska (ALDI Polska).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing_pt=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_spacing_pt:
        spacing.set(qn("w:line"), str(int(line_spacing_pt * 20)))
        spacing.set(qn("w:lineRule"), "exact")


def add_run(paragraph, text, size=11, bold=False, color="000000", font="Aptos"):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=12, after_pt=12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=6)
    add_run(p, text, size=16, bold=True, color="1F2937")


def add_section_header(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=14, after_pt=4)
    add_run(p, text, size=13, bold=True, color="1E40AF")


def add_subsection_header(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=10, after_pt=3)
    add_run(p, text, size=11, bold=True, color="374151")


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=1, after_pt=1)
    add_run(p, f"{label}: ", size=11, bold=True, color="4B5563")
    add_run(p, value, size=11, color="111827")


def add_body_text(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=2, after_pt=2, line_spacing_pt=15)
    add_run(p, text, size=11, color="111827")


def add_email_block(doc, label, subject, body):
    add_subsection_header(doc, label)
    add_label_value(doc, "Temat", subject)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=4, after_pt=4, line_spacing_pt=15)
    # Split body by newlines for proper paragraph handling
    lines = body.split("\n")
    first = True
    for line in lines:
        if first:
            add_run(p, line, size=11, color="111827")
            first = False
        else:
            # New paragraph for each line break
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before_pt=2, after_pt=2, line_spacing_pt=15)
            add_run(p, line, size=11, color="111827")


def add_quote_block(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=4, after_pt=4, line_spacing_pt=15)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    add_run(p, text, size=11, color="4B5563")


# ─────────────────────────────────────────────────────────────────────────────
# CAMPAIGN DATA
# ─────────────────────────────────────────────────────────────────────────────

DAMIAN_HUCZ = {
    "name": "Damian Hucz",
    "title": "Purchasing Director",
    "company": "AB Bechcicki Sp. z o.o.",
    "industry": "Dystrybucja materiałów budowlanych (hurtownia B2B)",
    "icp_fit": "Wysoki - Purchasing Director w firmie z 300+ dostawcami, dużym portfolio produktowym (100 tys.+ SKU) i wieloma kategoriami zakupowymi. Bezpośrednio odpowiada za negocjacje, warunki handlowe i koszty zakupu w skali firmy.",
    "tier": "Tier 2 - Procurement Management / Dyrektorzy zakupów",
    "tier_rationale": (
        "Damian Hucz jako Purchasing Director odpowiada za funkcję zakupową AB Bechcicki - firmy z ponad 300 dostawcami "
        "i portfelem obejmującym izolacje, chemię budowlaną, farby, płytki, suchą zabudowę, dachy i wiele innych kategorii. "
        "To klasyczny Tier 2: lider zakupów odpowiedzialny za savings delivery, jakość pracy zespołu, warunki handlowe "
        "z dostawcami i raportowanie efektów. Nie jest C-Level (Tier 1) ani operacyjnym kupcem (Tier 3)."
    ),
    "tensions": (
        "AB Bechcicki to duży dystrybutor budowlany z setkami dostawców - producentów materiałów budowlanych "
        "(Atlas, Rockwool, Siniat, Velux, Tikkurila, Ceresit, Kerakoll i wielu innych). "
        "Kluczowe napięcia:\n"
        "- Presja cenowa ze strony producentów materiałów budowlanych - ceny cementu, chemii, izolacji, opakowań rosną.\n"
        "- Przy 300+ dostawcach trudno utrzymać jednolity standard oceny podwyżek i przygotowania negocjacji.\n"
        "- Sezonowość budownictwa (szczyt wiosna-lato) oznacza, że negocjacje warunków na nowy sezon mają konkretny deadline.\n"
        "- Marża dystrybutora jest z natury niska - każdy punkt procentowy na warunkach zakupu przekłada się na wynik.\n"
        "- Producenci mają coraz silniejszą pozycję negocjacyjną (konsolidacja rynku materiałów budowlanych).\n"
        "- Savings delivery przy takiej skali portfela wymaga systemowego podejścia."
    ),
    "angle": (
        "Standard przygotowania negocjacji z dostawcami przy dużym portfelu kategorii budowlanych. "
        "Przy 300+ dostawcach i setkach tysięcy SKU kluczowe jest, żeby zespół zakupowy pracował "
        "według powtarzalnego standardu oceny podwyżek i przygotowania do rozmów - "
        "bo nierówna jakość przygotowania między kategoriami kosztuje."
    ),
    "hypothesis": (
        "Zakładam, że przy ponad 300 dostawcach materiałów budowlanych i szerokim portfelu kategorii - od chemii budowlanej "
        "i izolacji, przez farby, po suchą zabudowę - kluczowym pytaniem może być dziś to, jak zapewnić powtarzalny standard "
        "przygotowania negocjacji w całym zespole zakupowym. W dystrybucji budowlanej, gdzie marża jest wrażliwa "
        "na każdą zmianę warunków, nierówna jakość przygotowania między kategoriami bywa kosztowna."
    ),

    # Email 1
    "e1_subject": "Pytanie o przygotowanie negocjacji z dostawcami - AB Bechcicki",
    "e1_body": (
        "Dzień dobry Panie Damianie,\n\n"
        "zwróciłem uwagę, że AB Bechcicki pracuje z szerokim portfelem dostawców materiałów budowlanych - "
        "od producentów chemii i izolacji, przez farby, po suchą zabudowę i systemy dachowe.\n\n"
        "Przy takiej skali i liczbie kategorii domyślam się, że coraz ważniejsze staje się, żeby zespół zakupowy "
        "przygotowywał się do rozmów z dostawcami według powtarzalnego standardu - szczególnie gdy producenci "
        "uzasadniają podwyżki zmianami cen surowców, energii czy logistyki.\n\n"
        "Jestem Tomasz Uściński z Profitii. Pomagamy dyrektorom zakupów w firmach dystrybucyjnych systemowo "
        "poprawiać jakość przygotowania negocjacji z dostawcami, tak żeby wynik mniej zależał od tego, "
        "jak dobrze dany kupiec zna swoją kategorię.\n\n"
        "Jeśli temat jest dla Pana aktualny, proszę wybrać dogodny termin tutaj: [link do Calendly].\n"
        "Może Pan też po prostu odpisać \"TAK\" i podać numer telefonu - oddzwonię."
    ),

    # Follow-up 1
    "fu1_subject": "RE: Pytanie o przygotowanie negocjacji z dostawcami - AB Bechcicki",
    "fu1_body": (
        "Dzień dobry Panie Damianie,\n\n"
        "w kontekście dystrybucji materiałów budowlanych kluczowe bywa szybkie rozdzielenie, "
        "które podwyżki producentów wynikają z realnych zmian kosztowych, a które są próbą poprawienia marży dostawcy "
        "pod pretekstem wzrostu cen surowców czy energii.\n\n"
        "W takich sytuacjach pomagamy zbudować wspólny standard oceny podwyżek: jak weryfikować argumenty dostawcy, "
        "jak porównywać z cost driverami i benchmarkami rynkowymi, jak ocenić, gdzie warto wrócić do negocjacji "
        "z dostawcami warunków.\n\n"
        "Jeśli chce Pan, mogę pokazać na krótkiej rozmowie, jak taki standard przygotowania wygląda "
        "w praktyce dystrybucji budowlanej - tutaj można wybrać termin: [link do Calendly], "
        "albo po prostu odpisać \"TAK\" i podać numer telefonu."
    ),

    # Follow-up 2
    "fu2_subject": "RE: Pytanie o przygotowanie negocjacji z dostawcami - AB Bechcicki",
    "fu2_body": (
        "Dzień dobry Panie Damianie,\n\n"
        "z naszych obserwacji wynika, że w dystrybucji budowlanej najtrudniejszy moment pojawia się, "
        "gdy zarząd oczekuje poprawy warunków zakupowych, a zespół nie ma wspólnej logiki, żeby pokazać, "
        "które kategorie mają największy potencjał i gdzie avoided cost jest realny, a nie deklaratywny.\n\n"
        "Pomagamy uporządkować to na poziomie portfela kategorii - wskazać, gdzie wrócić do dostawców, "
        "jak przygotować argumentację i jak lepiej raportować efekt zakupowy do zarządu. Często POC "
        "na jednej kategorii wystarczy, żeby ocenić wartość podejścia.\n\n"
        "Jeśli chce Pan, możemy omówić to na przykładzie jednej z kategorii - tutaj można wybrać termin: "
        "[link do Calendly], albo po prostu odpisać \"TAK\" i podać numer telefonu."
    ),

    "rationale": (
        "Kampania jest relewantna, ponieważ:\n\n"
        "DLA OSOBY: Damian Hucz jako Purchasing Director odpowiada za wynik zakupowy całego portfela - "
        "a przy 300+ dostawcach materiałów budowlanych wyzwanie polega na zapewnieniu powtarzalnej jakości "
        "przygotowania negocjacji w wielu kategoriach jednocześnie. To klasyczny pain point Tier 2.\n\n"
        "DLA FIRMY: AB Bechcicki to dystrybutor z inherentnie niską marżą - każda poprawa warunków zakupowych "
        "przekłada się bezpośrednio na wynik. Przy szerokim portfelu (chemia, izolacje, farby, płytki, dachy) "
        "standaryzacja podejścia do negocjacji ma wymierny efekt.\n\n"
        "DLA BRANŻY: Dystrybucja materiałów budowlanych mierzy się z konsolidacją producentów, "
        "rosnącymi cenami surowców i energii, sezonowością. Producenci coraz częściej uzasadniają podwyżki "
        "zmianami kosztowymi - pytanie, ile z tych podwyżek jest realnie uzasadnionych.\n\n"
        "CO ODRÓŻNIA OD GENERYCZNEGO OUTREACHU: Kampania mówi o konkretnych realiach dystrybucji budowlanej "
        "(portfel kategorii, producenci materiałów, sezonowość), nie o \"zakupach\" ogólnie. "
        "Każdy mail buduje inny kąt - standard przygotowania, weryfikacja podwyżek, raportowanie do zarządu."
    ),
}

MONIKA_SITKOWSKA = {
    "name": "Monika Sitkowska",
    "title": "Senior Director Purchasing / Dyrektor ds. Zakupów",
    "company": "ALDI Polska",
    "industry": "Retail / Sieć dyskontowa (FMCG, private label)",
    "icp_fit": "Wysoki - Senior Director Purchasing w dużej sieci dyskontowej z szerokim portfelem dostawców FMCG, silnym naciskiem na koszty i cenę na półce. Bezpośrednio odpowiada za warunki zakupowe, negocjacje z dostawcami i efektywność kosztową.",
    "tier": "Tier 2 - Procurement Management / Dyrektorzy zakupów",
    "tier_rationale": (
        "Monika Sitkowska jako Senior Director Purchasing odpowiada za funkcję zakupową ALDI Polska - "
        "sieci dyskontowej, dla której efektywność kosztowa zakupów jest fundamentem modelu biznesowego. "
        "To Tier 2: lider zakupów odpowiedzialny za negocjacje z dostawcami, savings delivery, "
        "warunki handlowe i jakość pracy zespołu w skali organizacji. "
        "Seniority jest bardzo wysoka (Senior Director), ale rola jest procurement-management, nie C-Level."
    ),
    "tensions": (
        "ALDI Polska to sieć dyskontowa z modelem EDLP (Every Day Low Price), "
        "gdzie efektywność zakupowa jest DNA firmy. Kluczowe napięcia:\n"
        "- Model dyskontowy wymaga ekstremalnej dyscypliny kosztowej - każdy grosz na warunkach zakupu "
        "przekłada się na cenę na półce i przewagę konkurencyjną.\n"
        "- Dostawcy FMCG (producenci żywności, napojów, chemii, private label) regularnie uzasadniają podwyżki "
        "kosztami surowców, opakowań, energii i logistyki.\n"
        "- Private label to rdzeń strategii ALDI - a tam negocjacje z dostawcami są szczególnie intensywne, "
        "bo cena i specyfikacja są definiowane przez sieć.\n"
        "- Presja ze strony konkurencji (Lidl, Biedronka) wymaga ciągłej ochrony pozycji cenowej.\n"
        "- Skala operacji (setki sklepów, szerokie portfolio dostawców) wymaga systematyczności "
        "w ocenie warunków i podwyżek.\n"
        "- Zarząd oczekuje mierzalnych efektów zakupowych i obrony marży."
    ),
    "angle": (
        "Systematyczna weryfikacja zasadności podwyżek dostawców FMCG i private label w modelu dyskontowym. "
        "Gdy model biznesowy opiera się na najniższej cenie na półce, każda zaakceptowana podwyżka, "
        "która nie jest w pełni uzasadniona, osłabia pozycję konkurencyjną."
    ),
    "hypothesis": (
        "Zakładam, że w modelu dyskontowym, gdzie cena na półce jest kluczowym elementem przewagi, "
        "szczególnie ważna może być dziś systematyczna weryfikacja, które podwyżki dostawców FMCG "
        "i private label są realnie uzasadnione zmianami kosztowymi, a które wymagają twardszej rozmowy. "
        "Przy takiej skali operacji i presji konkurencyjnej nierówna jakość oceny podwyżek "
        "w portfelu kategorii może kosztować więcej, niż widać na poziomie pojedynczej negocjacji."
    ),

    # Email 1
    "e1_subject": "Pytanie o weryfikację podwyżek dostawców - ALDI Polska",
    "e1_body": (
        "Dzień dobry Pani Moniko,\n\n"
        "patrząc na dynamikę rynku FMCG i private label w Polsce, domyślam się, że jako Dyrektor ds. Zakupów "
        "w ALDI może Pani mierzyć się teraz z falą podwyżek od dostawców uzasadnianych kosztami surowców, "
        "opakowań i logistyki.\n\n"
        "W modelu dyskontowym, gdzie cena na półce jest fundamentem przewagi, kluczowe bywa szybkie oddzielenie "
        "podwyżek realnie wynikających ze zmian kosztowych od tych, które warto zakwestionować "
        "w rozmowach z dostawcami.\n\n"
        "Jestem Tomasz Uściński z Profitii. Pomagamy dyrektorom zakupów w sieciach handlowych ograniczać "
        "nieuzasadnione podwyżki dostawców FMCG, weryfikując je na poziomie cost driverów i realiów rynkowych.\n\n"
        "Jeśli temat jest dla Pani aktualny, proszę wybrać dogodny termin tutaj: [link do Calendly].\n"
        "Może Pani też po prostu odpisać \"TAK\" i podać numer telefonu - oddzwonię."
    ),

    # Follow-up 1
    "fu1_subject": "RE: Pytanie o weryfikację podwyżek dostawców - ALDI Polska",
    "fu1_body": (
        "Dzień dobry Pani Moniko,\n\n"
        "w kontekście negocjacji z dostawcami FMCG i private label najtrudniejsze bywa nie samo odrzucenie podwyżki, "
        "ale zbudowanie spójnego standardu oceny: jak systematycznie weryfikować, czy argumenty dostawcy "
        "(surowce, energia, opakowania, transport) odpowiadają temu, co faktycznie dzieje się na rynku.\n\n"
        "Pomagamy zespołom zakupowym w retail przełożyć dane o cost driverach i benchmarkach rynkowych "
        "na konkretną argumentację, która pozwala ograniczać nieuzasadnione podwyżki "
        "i lepiej przygotowywać zespół do rozmów z dostawcami.\n\n"
        "Jeśli chce Pani, mogę pokazać, jak taki standard przygotowania wygląda w praktyce - "
        "tutaj można wybrać termin: [link do Calendly], "
        "albo po prostu odpisać \"TAK\" i podać numer telefonu."
    ),

    # Follow-up 2
    "fu2_subject": "RE: Pytanie o weryfikację podwyżek dostawców - ALDI Polska",
    "fu2_body": (
        "Dzień dobry Pani Moniko,\n\n"
        "z naszych obserwacji wynika, że w sieciach handlowych punkt decyzyjny często pojawia się, "
        "gdy zarząd pyta o mierzalne efekty zakupowe - avoided cost, obronioną marżę, "
        "ograniczone podwyżki - a zespół nie ma jeszcze wspólnej logiki raportowania tych efektów.\n\n"
        "W takich sytuacjach pomagamy uporządkować to na poziomie portfela kategorii: "
        "wskazać, gdzie efekt jest realny, jak go pokazać i gdzie warto wrócić do rozmów z dostawcami. "
        "Często POC na jednej kategorii private label wystarczy, żeby ocenić wartość podejścia.\n\n"
        "Jeśli chce Pani, możemy omówić to na konkretnym przykładzie - tutaj można wybrać termin: "
        "[link do Calendly], albo po prostu odpisać \"TAK\" i podać numer telefonu."
    ),

    "rationale": (
        "Kampania jest relewantna, ponieważ:\n\n"
        "DLA OSOBY: Monika Sitkowska jako Senior Director Purchasing odpowiada za efektywność zakupową "
        "ALDI Polska. Przy jej seniority kluczowe jest nie samo negocjowanie, ale zapewnienie systemowego "
        "podejścia do weryfikacji podwyżek i raportowanie efektów zakupowych do zarządu.\n\n"
        "DLA FIRMY: ALDI to sieć dyskontowa, gdzie efektywność zakupowa jest fundamentem modelu biznesowego. "
        "Każda zaakceptowana podwyżka, która nie jest uzasadniona, osłabia pozycję cenową na półce "
        "i przewagę wobec Lidla i Biedronki.\n\n"
        "DLA BRANŻY: Retail FMCG w Polsce mierzy się z falą podwyżek dostawców uzasadnianych inflacją, "
        "cenami surowców i energii. Część z nich jest zasadna, część nie. Sieci dyskontowe "
        "z modelem EDLP są najbardziej wrażliwe na tę dynamikę.\n\n"
        "CO ODRÓŻNIA OD GENERYCZNEGO OUTREACHU: Kampania mówi o konkretnych realiach dyskontu "
        "(cena na półce, private label, model EDLP, presja konkurencyjna Lidl/Biedronka), "
        "nie o \"zakupach\" ogólnie. Każdy mail buduje inny kąt - weryfikacja podwyżek, "
        "standard oceny, raportowanie avoided cost do zarządu."
    ),
}


def build_campaign_doc(data, output_path):
    """Build a Word document for one campaign."""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    add_title(doc, f"Kampania outbound: {data['name']}")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=2)
    add_run(p, f"{data['title']} | {data['company']}", size=11, color="6B7280")

    add_horizontal_line(doc)

    # 1. Dane kampanii
    add_section_header(doc, "1. Dane kampanii")
    add_label_value(doc, "Imię i nazwisko", data["name"])
    add_label_value(doc, "Stanowisko", data["title"])
    add_label_value(doc, "Firma", data["company"])
    add_label_value(doc, "Branża", data["industry"])
    add_label_value(doc, "Nadawca", "Tomasz Uściński, Head of Sales, Profitia")

    # 2. ICP i Tier
    add_section_header(doc, "2. ICP i Tier")
    add_label_value(doc, "ICP fit", data["icp_fit"])
    add_label_value(doc, "Tier", data["tier"])
    add_subsection_header(doc, "Uzasadnienie Tieru")
    add_body_text(doc, data["tier_rationale"])

    # 3. Kontekst firmy i branży
    add_section_header(doc, "3. Kontekst firmy i branży")
    add_subsection_header(doc, "Główne napięcia biznesowe")
    add_body_text(doc, data["tensions"])
    add_subsection_header(doc, "Messaging angle")
    add_body_text(doc, data["angle"])

    # 4. Hipoteza biznesowa
    add_section_header(doc, "4. Hipoteza biznesowa")
    add_quote_block(doc, data["hypothesis"])

    # 5. Email 1
    add_section_header(doc, "5. Email 1")
    add_email_block(doc, "EMAIL 1", data["e1_subject"], data["e1_body"])

    # 6. Follow-up 1
    add_section_header(doc, "6. Follow-up 1 (step 2)")
    add_email_block(doc, "FOLLOW-UP 1", data["fu1_subject"], data["fu1_body"])

    # 7. Follow-up 2
    add_section_header(doc, "7. Follow-up 2 (step 3)")
    add_email_block(doc, "FOLLOW-UP 2", data["fu2_subject"], data["fu2_body"])

    # 8. Uzasadnienie kampanii
    add_section_header(doc, "8. Uzasadnienie kampanii")
    add_body_text(doc, data["rationale"])

    doc.save(output_path)
    print(f"  ✅ Zapisano: {output_path}")


def build_comparison_doc(output_path):
    """Build comparison summary document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title(doc, "Porównanie kampanii outbound")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=2)
    add_run(p, "Damian Hucz (AB Bechcicki) vs Monika Sitkowska (ALDI Polska)", size=11, color="6B7280")

    add_horizontal_line(doc)

    # Overview table-like section
    add_section_header(doc, "1. Porównanie profili")

    for label, val_d, val_m in [
        ("Stanowisko", "Purchasing Director", "Senior Director Purchasing"),
        ("Firma", "AB Bechcicki Sp. z o.o.", "ALDI Polska"),
        ("Branża", "Dystrybucja materiałów budowlanych (B2B)", "Retail dyskontowy (FMCG, private label)"),
        ("Tier", "Tier 2 - Procurement Management", "Tier 2 - Procurement Management"),
        ("Model biznesowy", "Hurtownia B2B, niska marża dystrybucyjna", "Dyskont EDLP, cena na półce jako fundament"),
        ("Portfel dostawców", "300+ producentów materiałów budowlanych", "Producenci FMCG, napojów, chemii, private label"),
        ("Główny pain point", "Standard przygotowania negocjacji przy dużym portfelu kategorii", "Systematyczna weryfikacja podwyżek dostawców FMCG"),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=2, after_pt=2)
        add_run(p, f"{label}:", size=11, bold=True, color="4B5563")
        p2 = doc.add_paragraph()
        set_paragraph_spacing(p2, before_pt=0, after_pt=0)
        add_run(p2, f"  Hucz: ", size=10, bold=True, color="1E40AF")
        add_run(p2, val_d, size=10, color="374151")
        p3 = doc.add_paragraph()
        set_paragraph_spacing(p3, before_pt=0, after_pt=2)
        add_run(p3, f"  Sitkowska: ", size=10, bold=True, color="1E40AF")
        add_run(p3, val_m, size=10, color="374151")

    add_horizontal_line(doc)

    # Differences
    add_section_header(doc, "2. Kluczowe różnice między kampaniami")

    differences = [
        (
            "Kontekst branżowy",
            "AB Bechcicki działa w dystrybucji B2B materiałów budowlanych - kupuje od producentów (Atlas, Rockwool, Siniat, Velux) "
            "i sprzedaje firmom budowlanym. Marża dystrybutora jest z natury niska, sezonowość wysoka. "
            "ALDI Polska to sieć dyskontowa FMCG z modelem EDLP - kupuje od producentów żywności, napojów i chemii, "
            "sprzedaje konsumentom. Efektywność zakupowa jest DNA modelu biznesowego."
        ),
        (
            "Messaging angle",
            "Hucz: Standard przygotowania negocjacji przy szerokim portfelu kategorii budowlanych - "
            "jak zapewnić powtarzalną jakość, gdy kategorie są bardzo różne (chemia, izolacje, farby, dachy). "
            "Sitkowska: Systematyczna weryfikacja podwyżek dostawców FMCG/private label - "
            "jak oddzielić uzasadnione wzrosty od narracji dostawców, gdy cena na półce jest fundamentem modelu."
        ),
        (
            "Sekwencja follow-upów",
            "Hucz: Email 1 = standard przygotowania → FU1 = weryfikacja podwyżek producentów → FU2 = avoided cost i raportowanie do zarządu + POC. "
            "Sitkowska: Email 1 = weryfikacja podwyżek FMCG → FU1 = standard oceny (cost drivers, benchmarki) → FU2 = raportowanie avoided cost + POC na private label."
        ),
        (
            "Napięcia specyficzne",
            "Hucz: Konsolidacja producentów budowlanych wzmacnia ich pozycję, sezonowość wymusza timing negocjacji, "
            "100 tys.+ SKU = ogromna złożoność portfela. "
            "Sitkowska: Presja konkurencyjna Lidl/Biedronka, private label jako rdzeń strategii, "
            "model EDLP wymaga ekstremalnej dyscypliny kosztowej."
        ),
    ]

    for title, desc in differences:
        add_subsection_header(doc, title)
        add_body_text(doc, desc)

    add_horizontal_line(doc)

    # Why each campaign is well-fitted
    add_section_header(doc, "3. Dlaczego każda kampania jest dobrze dopasowana")

    add_subsection_header(doc, "Damian Hucz - AB Bechcicki")
    add_body_text(
        doc,
        "Kampania mówi językiem dystrybucji materiałów budowlanych - nie ogólnym językiem \"zakupów\". "
        "Odnosi się do konkretnych kategorii (chemia budowlana, izolacje, farby, dachy), konkretnych typów dostawców "
        "(producenci materiałów budowlanych), konkretnych realiów branżowych (sezonowość, konsolidacja). "
        "Messaging angle (standard przygotowania przy dużym portfelu) trafia w główny pain point "
        "Purchasing Directora firmy dystrybucyjnej: jak systemowo dowozić wynik zakupowy przy setkach dostawców."
    )

    add_subsection_header(doc, "Monika Sitkowska - ALDI Polska")
    add_body_text(
        doc,
        "Kampania mówi językiem dyskontu FMCG - nie ogólnym językiem \"sieci handlowych\". "
        "Odnosi się do modelu EDLP, ceny na półce jako fundamentu, private label jako rdzenia strategii, "
        "presji konkurencyjnej (Lidl, Biedronka). Messaging angle (weryfikacja podwyżek FMCG) trafia "
        "w główny pain point Senior Director Purchasing w dyskoncie: jak ograniczać podwyżki, "
        "które osłabiają pozycję cenową - bo w dyskoncie każda niezasadna podwyżka to utrata przewagi."
    )

    add_horizontal_line(doc)

    # Final note
    add_section_header(doc, "4. Podsumowanie")
    add_body_text(
        doc,
        "Obie kampanie są Tier 2, ale różnią się fundamentalnie kontekstem branżowym, "
        "messaging angle'em i sekwencją follow-upów. Nie ma ryzyka, że odbiorca pomyśli "
        "\"ten sam mail dostaje każdy dyrektor zakupów\". "
        "Każda kampania jest osadzona w realiach konkretnej firmy, branży i roli - "
        "i brzmi jak przygotowana dla konkretnego człowieka, nie dla segmentu."
    )

    doc.save(output_path)
    print(f"  ✅ Zapisano: {output_path}")


def main():
    output_dir = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..", "outputs", "word_campaigns"
    )
    os.makedirs(output_dir, exist_ok=True)

    print("Generuję kampanie outbound...\n")

    # Campaign 1
    path1 = os.path.join(output_dir, "Damian_Hucz_campaign.docx")
    build_campaign_doc(DAMIAN_HUCZ, path1)

    # Campaign 2
    path2 = os.path.join(output_dir, "Monika_Sitkowska_campaign.docx")
    build_campaign_doc(MONIKA_SITKOWSKA, path2)

    # Comparison
    path3 = os.path.join(output_dir, "campaign_summary_comparison.docx")
    build_comparison_doc(path3)

    print(f"\nGotowe. Pliki w: {output_dir}")


if __name__ == "__main__":
    main()
