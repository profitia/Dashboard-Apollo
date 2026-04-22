#!/usr/bin/env python3
"""
Generate Tier 2 Real Pilot outputs:
- tests/output/tier2_real_pilot_selected_leads.json
- tests/output/tier2_real_pilot_campaigns.json
- tests/output/tier2_real_pilot_selection_report.md
- outputs/word_campaigns/tier2_real_pilot_review.docx
"""

import json
import os
from datetime import datetime

# ============================================================
# LEAD DATA — Selected 3 from CSV
# ============================================================

ALL_LEADS_COUNT = 9

selected_leads = [
    {
        "id": 1,
        "first_name": "Konrad",
        "last_name": "Ludwicki",
        "vocative": "Konradzie",
        "gender": "male",
        "title": "Senior Director, Head of Procurement Europe",
        "company": "Circle K",
        "email": "kludwicki@circlek.com",
        "seniority": "Head",
        "departments": "Finance",
        "sub_departments": "Sourcing / Procurement",
        "industry": "retail",
        "city": "Piaseczno",
        "country": "Poland",
        "company_city": "Laval",
        "company_country": "Canada",
        "linkedin": "http://www.linkedin.com/in/konrad-l-42a00521",
        "website": "https://circlek.com",
        "keywords_summary": "retail, fuel, convenience, food & beverage, supply chain, sustainability, digital payments, franchise, store operations, fleet management",
        "selection_reason": (
            "Najsilniejszy lead w całym CSV. Head of Procurement Europe w Circle K "
            "to rola paneuropejska z odpowiedzialnością za wiele rynków i kategorii "
            "jednocześnie (paliwa, convenience, F&B, wyposażenie, usługi). Circle K "
            "to duży, rozpoznawalny brand z globalną skalą. Doskonałe dopasowanie do "
            "Tier 2 - lider funkcji zakupowej rozliczany z savings delivery w skali "
            "organizacji. Bogaty profil pozwala na trafną personalizację."
        ),
        "tier": "Tier 2",
        "persona": "Dyrektor Zakupów / Head of Procurement",
        "icp_fit": "Excellent - procurement leadership z paneuropejskim scope",
        "angle": "savings_delivery",
        "hypothesis": (
            "Przy odpowiedzialności za zakupy w wielu europejskich rynkach i zróżnicowanym "
            "portfelu kategorii, wyzwaniem jest spójne dowożenie oszczędności we wszystkich "
            "kategoriach jednocześnie, niezależnie od rynku i zespołu."
        ),
        "risk_assessment": "Niskie ryzyko. Profil idealnie pasuje do Tier 2. Bogata baza do personalizacji.",
    },
    {
        "id": 2,
        "first_name": "Anna",
        "last_name": "Polak",
        "vocative": "Anno",
        "gender": "female",
        "title": "Procurement Director Poland",
        "company": "DS Smith",
        "email": "anna.polak@dssmith.com",
        "seniority": "Director",
        "departments": "Finance",
        "sub_departments": "Sourcing / Procurement",
        "industry": "mechanical or industrial engineering / packaging & containers",
        "city": "Kielce",
        "country": "Poland",
        "company_city": "London",
        "company_country": "United Kingdom",
        "linkedin": "http://www.linkedin.com/in/anna-polak-501651106",
        "website": "https://dssmith.com",
        "keywords_summary": "packaging, corrugated, paper, recycling, sustainability, circular economy, supply chain, e-commerce packaging, industrial packaging",
        "selection_reason": (
            "Procurement Director Poland w DS Smith - dużej, giełdowej firmie opakowaniowej "
            "z globalnym zasięgiem. Branża opakowaniowa/papierowa jest silnie narażona na "
            "zmienność cen surowców (recyclat, celuloza, energia), co daje naturalny punkt "
            "wejścia do rozmowy o weryfikacji podwyżek dostawców. DS Smith to rozpoznawalna "
            "marka, co ułatwia budowanie wiarygodnej hipotezy. Bogaty profil danych."
        ),
        "tier": "Tier 2",
        "persona": "Dyrektor Zakupów / Procurement Director",
        "icp_fit": "Excellent - procurement director w dużej firmie produkcyjnej z branży wrażliwej na koszty surowców",
        "angle": "supplier_price_increases",
        "hypothesis": (
            "W branży opakowaniowej, gdzie ceny surowców papierowych, recyclatu i energii "
            "zmieniają się dynamicznie, kluczowym wyzwaniem jest ocena, które podwyżki "
            "dostawców są realnie uzasadnione kosztowo, a które są narracją negocjacyjną."
        ),
        "risk_assessment": "Niskie ryzyko. Silne dopasowanie branżowe. Naturalny pain point w opakowaniach.",
    },
    {
        "id": 3,
        "first_name": "Sebastian",
        "last_name": "Kacperek",
        "vocative": "Sebastianie",
        "gender": "male",
        "title": "Associate Director Global Procurement",
        "company": "AstraZeneca",
        "email": "sebastian.kacperek@astrazeneca.com",
        "seniority": "Director",
        "departments": "Finance",
        "sub_departments": "Sourcing / Procurement",
        "industry": "pharmaceuticals",
        "city": "Warsaw",
        "country": "Poland",
        "company_city": "Cambridge",
        "company_country": "United Kingdom",
        "linkedin": "http://www.linkedin.com/in/sebastiankacperek",
        "website": "https://astrazeneca.com",
        "keywords_summary": "pharmaceutical, biopharmaceutical, innovation, R&D, biotechnology, clinical trials, oncology, immunology, global healthcare",
        "selection_reason": (
            "Associate Director Global Procurement w AstraZeneca - jednej z największych "
            "firm farmaceutycznych na świecie. Farmacja to branża z zaawansowaną organizacją "
            "zakupową, gdzie wymagania jakościowe i regulacyjne zawężają pole negocjacji "
            "z dostawcami - co czyni temat przygotowania negocjacji szczególnie trafnym. "
            "Globalny scope roli daje podstawę do rozmowy o standaryzacji podejścia "
            "w wielu kategoriach. AstraZeneca jest powszechnie znana."
        ),
        "tier": "Tier 2",
        "persona": "Dyrektor Zakupów / Associate Director Global Procurement",
        "icp_fit": "Excellent - global procurement leadership w dużej organizacji farmaceutycznej",
        "angle": "negotiation_preparation",
        "hypothesis": (
            "W farmacji wymagania jakościowe i regulacyjne istotnie zawężają pole negocjacji "
            "z dostawcami. Wyzwaniem jest identyfikacja kategorii, w których mimo tych "
            "ograniczeń nadal jest realna przestrzeń do poprawy warunków handlowych."
        ),
        "risk_assessment": "Niskie ryzyko. AstraZeneca to top-tier firma. Hipoteza specyficzna dla farmacji.",
    },
]

# ============================================================
# CAMPAIGN CONTENT
# ============================================================

campaigns = {
    "konrad_ludwicki_circle_k": {
        "lead_id": 1,
        "contact": {
            "first_name": "Konrad",
            "last_name": "Ludwicki",
            "vocative": "Konradzie",
            "gender": "male",
            "title": "Senior Director, Head of Procurement Europe",
            "company": "Circle K",
            "email": "kludwicki@circlek.com",
        },
        "campaign_metadata": {
            "tier": "Tier 2",
            "angle": "savings_delivery",
            "persona": "Dyrektor Zakupów / Head of Procurement",
            "language": "pl",
            "campaign_type": "standard_outbound_tier2",
        },
        "email_1": {
            "subject": "Oszczędności w zakupach europejskich - Circle K",
            "body": (
                "Dzień dobry Panie Konradzie,\n\n"
                "zwróciłem uwagę na skalę operacji Circle K w Europie i Pana rolę jako "
                "Head of Procurement Europe. Przy tak zróżnicowanym portfelu kategorii "
                "- od convenience, przez F&B, po usługi i wyposażenie - domyślam się, "
                "że jednym z wyzwań jest spójne dowożenie oszczędności we wszystkich "
                "kategoriach jednocześnie.\n\n"
                "Nazywam się Tomasz Uściński, jestem z polskiej firmy Profitia, w której "
                "od 15 lat pomagamy dyrektorom zakupów w sieciach handlowych systemowo "
                "dowozić savings w wielu kategoriach - tak aby wynik nie zależał od stylu "
                "pracy pojedynczych kupców.\n\n"
                "Jeśli temat jest dla Pana aktualny, proszę wybrać dogodny termin tutaj: "
                "[link do Calendly].\n"
                'Może Pan też po prostu odpisać \u201eTAK\u201d i podać numer telefonu - oddzwonię.'
            ),
        },
        "follow_up_1": {
            "subject": "Re: Oszczędności w zakupach europejskich - Circle K",
            "body": (
                "Dzień dobry Panie Konradzie,\n\n"
                "dopowiem jedną rzecz w kontekście zakupów europejskich. Przy wielu "
                "rynkach i zespołach jednocześnie kluczowe jest, aby standard oceny "
                "ofert i podwyżek dostawców był porównywalny - niezależnie od kupca "
                "i kategorii.\n\n"
                "Pomagamy to osiągnąć, budując wspólną logikę przygotowania - od struktury "
                "kosztu, przez benchmarki, po argumentację negocjacyjną - tak aby każdy "
                "zespół startował z tego samego punktu w rozmowach z dostawcami.\n\n"
                "Jeśli chce Pan zobaczyć, jak to wygląda w praktyce, proszę dać znać "
                "- chętnie pokażę na jednej kategorii."
            ),
        },
        "follow_up_2": {
            "subject": "Re: Oszczędności w zakupach europejskich - Circle K",
            "body": (
                "Dzień dobry Panie Konradzie,\n\n"
                "z naszych obserwacji wynika, że dyrektorzy zakupów w dużych sieciach "
                "coraz częściej mierzą się z pytaniem: jak raportować uniknięte podwyżki "
                "jako realny wkład zakupów w wynik firmy?\n\n"
                "Pomagamy przełożyć dane o avoided cost na konkretne liczby do rozmowy "
                "z zarządem - tak aby efekty pracy zespołu zakupowego były widoczne "
                "i mierzalne.\n\n"
                "Jeśli temat jest aktualny, wystarczy krótka odpowiedź - umówimy się "
                "na 15 minut."
            ),
        },
        "apollo_custom_fields": {},
    },
    "anna_polak_ds_smith": {
        "lead_id": 2,
        "contact": {
            "first_name": "Anna",
            "last_name": "Polak",
            "vocative": "Anno",
            "gender": "female",
            "title": "Procurement Director Poland",
            "company": "DS Smith",
            "email": "anna.polak@dssmith.com",
        },
        "campaign_metadata": {
            "tier": "Tier 2",
            "angle": "supplier_price_increases",
            "persona": "Dyrektor Zakupów / Procurement Director",
            "language": "pl",
            "campaign_type": "standard_outbound_tier2",
        },
        "email_1": {
            "subject": "Weryfikacja podwyżek dostawców - DS Smith",
            "body": (
                "Dzień dobry Pani Anno,\n\n"
                "patrząc na sytuację w branży opakowaniowej - gdzie ceny surowców "
                "papierowych, recyclatu i energii zmieniają się dynamicznie z kwartału "
                "na kwartał - domyślam się, że jednym z kluczowych wyzwań w Pani roli "
                "jako Procurement Director jest ocena, które podwyżki dostawców są realnie "
                "uzasadnione kosztowo, a które są przede wszystkim narracją negocjacyjną.\n\n"
                "Nazywam się Tomasz Uściński, jestem z polskiej firmy Profitia, w której "
                "od 15 lat pomagamy zespołom zakupowym w firmach produkcyjnych weryfikować "
                "zasadność argumentów kosztowych dostawców - tak aby ograniczać "
                "nieuzasadnione podwyżki i chronić budżet zakupowy.\n\n"
                "Jeśli temat jest dla Pani aktualny, proszę wybrać dogodny termin tutaj: "
                "[link do Calendly].\n"
                'Może Pani też po prostu odpisać \u201eTAK\u201d i podać numer telefonu - oddzwonię.'
            ),
        },
        "follow_up_1": {
            "subject": "Re: Weryfikacja podwyżek dostawców - DS Smith",
            "body": (
                "Dzień dobry Pani Anno,\n\n"
                "dopowiem jedną rzecz w kontekście kosztów w branży opakowaniowej. Kiedy "
                "dostawca uzasadnia podwyżkę cenami recyclatu czy energii, kluczowe jest "
                "sprawdzenie, w jakim stopniu te zmiany faktycznie wpływają na koszt "
                "produktu - a w jakim to argument negocjacyjny.\n\n"
                "Pomagamy rozbić strukturę kosztu kategorii na konkretne drivery i porównać "
                "dynamikę cen z faktycznymi zmianami kosztowymi - tak aby rozmowa "
                "z dostawcą opierała się na danych, nie na deklaracjach.\n\n"
                "Jeśli chce Pani zobaczyć, jak to wygląda w praktyce na jednej kategorii, "
                "proszę dać znać."
            ),
        },
        "follow_up_2": {
            "subject": "Re: Weryfikacja podwyżek dostawców - DS Smith",
            "body": (
                "Dzień dobry Pani Anno,\n\n"
                "z perspektywy negocjacji z dostawcami często kluczowe jest nie tylko to, "
                "co negocjować, ale kiedy. W branży papierowej i opakowaniowej ceny "
                "surowców mają wyraźne cykle - i odpowiedni timing renegocjacji może mieć "
                "istotny wpływ na warunki.\n\n"
                "Pomagamy ocenić, czy aktualny moment jest dobry do rozmowy z dostawcą "
                "o cenie, indeksacji lub warunkach kontraktu - na podstawie prognoz "
                "i trendów kosztowych.\n\n"
                "Wystarczy krótka odpowiedź - umówimy się na 15 minut."
            ),
        },
        "apollo_custom_fields": {},
    },
    "sebastian_kacperek_astrazeneca": {
        "lead_id": 3,
        "contact": {
            "first_name": "Sebastian",
            "last_name": "Kacperek",
            "vocative": "Sebastianie",
            "gender": "male",
            "title": "Associate Director Global Procurement",
            "company": "AstraZeneca",
            "email": "sebastian.kacperek@astrazeneca.com",
        },
        "campaign_metadata": {
            "tier": "Tier 2",
            "angle": "negotiation_preparation",
            "persona": "Dyrektor Zakupów / Associate Director Global Procurement",
            "language": "pl",
            "campaign_type": "standard_outbound_tier2",
        },
        "email_1": {
            "subject": "Przygotowanie negocjacji z dostawcami - AstraZeneca",
            "body": (
                "Dzień dobry Panie Sebastianie,\n\n"
                "patrząc na specyfikę zakupów w branży farmaceutycznej - gdzie wymagania "
                "jakościowe i regulacyjne istotnie zawężają pole negocjacji z dostawcami "
                "- domyślam się, że jednym z wyzwań w Pana roli jest identyfikacja "
                "kategorii, w których mimo tych ograniczeń nadal jest realna przestrzeń "
                "do poprawy warunków handlowych.\n\n"
                "Nazywam się Tomasz Uściński, jestem z polskiej firmy Profitia, w której "
                "od 15 lat pomagamy zespołom zakupowym w dużych organizacjach skuteczniej "
                "przygotowywać się do negocjacji z dostawcami - nawet w kategoriach, "
                "gdzie wybór dostawców jest ograniczony.\n\n"
                "Jeśli temat jest dla Pana aktualny, proszę wybrać dogodny termin tutaj: "
                "[link do Calendly].\n"
                'Może Pan też po prostu odpisać \u201eTAK\u201d i podać numer telefonu - oddzwonię.'
            ),
        },
        "follow_up_1": {
            "subject": "Re: Przygotowanie negocjacji z dostawcami - AstraZeneca",
            "body": (
                "Dzień dobry Panie Sebastianie,\n\n"
                "dopowiem jedną rzecz w kontekście negocjacji z dostawcami w farmacji. "
                "Nawet przy ograniczonej bazie dostawców kluczowe pytanie często dotyczy "
                "nie samej ceny, ale konstrukcji kontraktu - indeksacji, wolumenu, "
                "warunków płatności czy długości umowy.\n\n"
                "Pomagamy sprawdzić, które z tych dźwigni mogą realnie przełożyć się "
                "na niższy koszt zakupu lub lepsze warunki handlowe - bez naruszania "
                "standardów jakościowych i compliance.\n\n"
                "Jeśli chce Pan zobaczyć to na przykładzie jednej kategorii, proszę "
                "dać znać."
            ),
        },
        "follow_up_2": {
            "subject": "Re: Przygotowanie negocjacji z dostawcami - AstraZeneca",
            "body": (
                "Dzień dobry Panie Sebastianie,\n\n"
                "w dużych organizacjach farmaceutycznych dyrektorzy zakupów często mierzą "
                "się z pytaniem: jak wykazać wartość unikniętych podwyżek jako realny "
                "wkład zakupów - szczególnie w kategoriach, gdzie oszczędności cenowe "
                "są trudniejsze do uzyskania?\n\n"
                "Pomagamy przełożyć efekty pracy zespołu zakupowego na konkretne dane "
                "do raportowania - tak aby wkład zakupów w wynik firmy był widoczny "
                "dla zarządu.\n\n"
                "Wystarczy krótka odpowiedź - umówimy się na 15 minut."
            ),
        },
        "apollo_custom_fields": {},
    },
}

# Build apollo_custom_fields for each campaign
for key, camp in campaigns.items():
    camp["apollo_custom_fields"] = {
        "sg_email_step_1_subject": camp["email_1"]["subject"],
        "sg_email_step_1_body": camp["email_1"]["body"],
        "sg_email_step_2_subject": camp["follow_up_1"]["subject"],
        "sg_email_step_2_body": camp["follow_up_1"]["body"],
        "sg_email_step_3_subject": camp["follow_up_2"]["subject"],
        "sg_email_step_3_body": camp["follow_up_2"]["body"],
    }


def count_words(text):
    """Count words in text, excluding formatting."""
    import re
    clean = re.sub(r'[„"""]', '', text)
    clean = re.sub(r'\[.*?\]', '', clean)
    return len(clean.split())


# ============================================================
# OUTPUT 1: selected_leads.json
# ============================================================

def generate_selected_leads():
    output = {
        "generated_at": datetime.now().isoformat(),
        "source_csv": "CSV do kampanii/2026-04-20 Tier 2 do VSC - test 10.csv",
        "total_leads_in_csv": ALL_LEADS_COUNT,
        "selected_count": 3,
        "selection_criteria": [
            "Dopasowanie do Tier 2 (procurement management / leadership)",
            "Jakość stanowiska (Director / Head / Senior Procurement leadership)",
            "Jakość i rozpoznawalność firmy",
            "Dopasowanie branży do logiki SpendGuru / Profitia",
            "Potencjał do personalizacji (bogactwo danych wejściowych)",
            "Możliwość zbudowania wiarygodnej hipotezy biznesowej",
        ],
        "selected_leads": selected_leads,
    }
    return output


# ============================================================
# OUTPUT 2: campaigns.json
# ============================================================

def generate_campaigns():
    output = {
        "generated_at": datetime.now().isoformat(),
        "campaign_type": "standard_outbound_tier2",
        "pilot_name": "Tier 2 Real Pilot - First 3",
        "sequence_cadence": "D0 / D+2 / D+2",
        "apollo_sync_status": "prepared_not_synced",
        "note": "Kampanie gotowe do review. Sync do Apollo po zatwierdzeniu.",
        "calendly_placeholder": "[link do Calendly] - wstaw rzeczywisty link przed synciem",
        "campaigns": campaigns,
    }
    return output


# ============================================================
# OUTPUT 3: selection_report.md
# ============================================================

def generate_report():
    lines = []
    lines.append("# Tier 2 Real Pilot - Raport selekcji\n")
    lines.append(f"**Data:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    lines.append(f"**Źródło:** `CSV do kampanii/2026-04-20 Tier 2 do VSC - test 10.csv`\n")
    lines.append("---\n")

    # 1. Summary
    lines.append("## 1. Podsumowanie CSV\n")
    lines.append(f"- Liczba leadów w CSV: **{ALL_LEADS_COUNT}**")
    lines.append("- Wybrano do pilota: **3**")
    lines.append("- Typ kampanii: **Standard Outbound Tier 2**")
    lines.append("- Sekwencja: **3 kroki (Email 1 + FU1 + FU2)**")
    lines.append("- Cadence: **D0 / D+2 / D+2**\n")

    # 2. Selected leads
    lines.append("## 2. Wybrane osoby\n")
    lines.append("| # | Osoba | Stanowisko | Firma | Branża | Angle |")
    lines.append("|---|-------|-----------|-------|--------|-------|")
    for lead in selected_leads:
        lines.append(
            f"| {lead['id']} | {lead['first_name']} {lead['last_name']} | "
            f"{lead['title']} | {lead['company']} | {lead['industry']} | "
            f"{lead['angle']} |"
        )
    lines.append("")

    # 3. Why these 3
    lines.append("## 3. Dlaczego te 3 osoby\n")
    for lead in selected_leads:
        lines.append(f"### {lead['first_name']} {lead['last_name']} - {lead['company']}\n")
        lines.append(f"**Stanowisko:** {lead['title']}")
        lines.append(f"**Tier:** {lead['tier']}")
        lines.append(f"**Persona:** {lead['persona']}")
        lines.append(f"**ICP Fit:** {lead['icp_fit']}")
        lines.append(f"**Angle:** {lead['angle']}")
        lines.append(f"\n**Powód wyboru:**\n{lead['selection_reason']}\n")
        lines.append(f"**Hipoteza:**\n{lead['hypothesis']}\n")
        lines.append(f"**Ocena ryzyka:** {lead['risk_assessment']}\n")

    # 4. Leads NOT selected
    lines.append("## 4. Osoby nieuwzględnione i powody\n")
    rejected = [
        ("Piotr Makal", "El-Cab", "Dobry profil, ale mniejsza firma z mniejszą rozpoznawalnością. Mniej danych do personalizacji niż top 3."),
        ("Agnieszka Świtón", "Gala Poland / Korona Candles", "Director Global Procurement - dobra rola, ale bardzo ograniczone dane w CSV (sparse keywords). Utrudnia budowanie hipotezy."),
        ("Michał Kulbaka", "Santander Consumer Bank", "Bankowość to słabsze dopasowanie do SpendGuru (negotiation-first, procurement-oriented). Procurement w banku to głównie IT i usługi - mniej trafne."),
        ("Piotr Torbé", "ROHLIG SUUS Logistics", "Zbyt wyspecjalizowana rola (seafreight FCL procurement). Nie typowy procurement leadership."),
        ("Tomasz Studniarek", "Neonet", "Mieszana rola (IT / Supply Chain / Procurement / Back Office). Nie jest czystym liderem zakupów."),
        ("Dariusz Strzemieczny", "ID Logistics", "Procurement Director - dobra rola, ale bardzo ograniczone dane w CSV. Logistyka jako branża mniej idealna."),
    ]
    for name, company, reason in rejected:
        lines.append(f"- **{name}** ({company}): {reason}")
    lines.append("")

    # 5. Quality assessment
    lines.append("## 5. Ocena jakości personalizacji\n")
    lines.append("| Lead | Jakość personalizacji | Ryzyko |")
    lines.append("|------|----------------------|--------|")
    for lead in selected_leads:
        lines.append(
            f"| {lead['first_name']} {lead['last_name']} ({lead['company']}) | "
            f"Wysoka | Niskie |"
        )
    lines.append("")
    lines.append("**Wszystkie 3 kampanie:**")
    lines.append("- Osadzone w realiach firmy i branży odbiorcy")
    lines.append("- Hipoteza specyficzna, nie generyczna")
    lines.append("- Angle dopasowany do kontekstu")
    lines.append("- Ton Tier 2 (procurement management, savings delivery, standard pracy zespołu)")
    lines.append("- Follow-upy wnoszą nową wartość (różne kąty)")
    lines.append("- CTA Tier 2 z Calendly + alternatywa TAK + telefon")
    lines.append("- Brak em dash, poprawna typografia PL")
    lines.append("- Brak podpisu (dodawany automatycznie)\n")

    # 6. Angles
    lines.append("## 6. Angle per osoba\n")
    lines.append("| Lead | Angle E1 | Kąt FU1 | Kąt FU2 |")
    lines.append("|------|----------|---------|---------|")
    lines.append("| Konrad Ludwicki | savings_delivery | team_standardization | avoided_cost_reporting |")
    lines.append("| Anna Polak | supplier_price_increases | cost_driver_analysis | timing_renegotiation |")
    lines.append("| Sebastian Kacperek | negotiation_preparation | contract_levers | avoided_cost_reporting |")
    lines.append("")

    # 7. Apollo readiness
    lines.append("## 7. Gotowość do Apollo\n")
    lines.append("- Custom fields (sg_email_step_1-3_subject/body): **przygotowane**")
    lines.append("- Calendly link: **placeholder [link do Calendly] - wymaga wstawienia przed synciem**")
    lines.append("- Cadence D0/D+2/D+2: **zgodny z centralną konfiguracją**")
    lines.append("- Sync do Apollo: **NIE wykonany - czeka na zatwierdzenie**\n")

    # 8. Recommendation
    lines.append("## 8. Rekomendacja\n")
    lines.append("### Konrad Ludwicki (Circle K)")
    lines.append("**GOTOWE DO PILOTA**")
    lines.append("Kampania silnie spersonalizowana, angle trafny, hipoteza wiarygodna. Jedyne co wymaga uzupełnienia: link Calendly.\n")
    lines.append("### Anna Polak (DS Smith)")
    lines.append("**GOTOWE DO PILOTA**")
    lines.append("Naturalne osadzenie w branży opakowaniowej. Podwyżki surowców to żywy temat w tej branży. Gotowe po wstawieniu linku Calendly.\n")
    lines.append("### Sebastian Kacperek (AstraZeneca)")
    lines.append("**GOTOWE DO PILOTA**")
    lines.append("Specyfika farmacji dobrze oddana. Hipoteza o ograniczeniach regulacyjnych = naturalna i trafna. Gotowe po wstawieniu linku Calendly.\n")

    lines.append("---\n")
    lines.append("**Ogólna rekomendacja:** Wszystkie 3 kampanie gotowe do pilota po wstawieniu linku Calendly.")
    lines.append("Przed wysyłką zalecany jest review treści przez nadawcę (Tomasz Uściński).")

    return "\n".join(lines)


# ============================================================
# OUTPUT 4: Word document
# ============================================================

def generate_word():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # Title
    title = doc.add_heading('Tier 2 Real Pilot - Review kampanii', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Wygenerowano: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Źródło: CSV do kampanii/2026-04-20 Tier 2 do VSC - test 10.csv')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Summary table
    doc.add_heading('Podsumowanie selekcji', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    cells = [
        ("Leadów w CSV", str(ALL_LEADS_COUNT)),
        ("Wybranych do pilota", "3"),
        ("Typ kampanii", "Standard Outbound Tier 2"),
        ("Cadence", "D0 / D+2 / D+2"),
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_page_break()

    # For each lead
    campaign_keys = [
        "konrad_ludwicki_circle_k",
        "anna_polak_ds_smith",
        "sebastian_kacperek_astrazeneca",
    ]

    for idx, (lead, camp_key) in enumerate(zip(selected_leads, campaign_keys)):
        camp = campaigns[camp_key]

        # Lead header
        doc.add_heading(
            f'Lead {idx + 1}: {lead["first_name"]} {lead["last_name"]} - {lead["company"]}',
            level=1,
        )

        # 1. Dane osoby i firmy
        doc.add_heading('Dane osoby i firmy', level=2)
        info_table = doc.add_table(rows=9, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_data = [
            ("Imię i nazwisko", f'{lead["first_name"]} {lead["last_name"]}'),
            ("Stanowisko", lead["title"]),
            ("Firma", lead["company"]),
            ("Email", lead["email"]),
            ("Branża", lead["industry"]),
            ("Lokalizacja", f'{lead["city"]}, {lead["country"]}'),
            ("Firma HQ", f'{lead["company_city"]}, {lead["company_country"]}'),
            ("LinkedIn", lead["linkedin"]),
            ("Website", lead["website"]),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value

        # 2. Dlaczego wybrany
        doc.add_heading('Dlaczego wybrany', level=2)
        doc.add_paragraph(lead["selection_reason"])

        # 3. ICP / Tier
        doc.add_heading('ICP / Tier', level=2)
        doc.add_paragraph(f'Tier: {lead["tier"]}')
        doc.add_paragraph(f'Persona: {lead["persona"]}')
        doc.add_paragraph(f'ICP Fit: {lead["icp_fit"]}')

        # 4. Angle
        doc.add_heading('Angle', level=2)
        doc.add_paragraph(lead["angle"])

        # 5. Hypothesis
        doc.add_heading('Hypothesis', level=2)
        doc.add_paragraph(lead["hypothesis"])

        # 6. Email 1
        doc.add_heading('Email 1', level=2)

        # Subject
        p = doc.add_paragraph()
        run = p.add_run('Temat: ')
        run.bold = True
        p.add_run(camp["email_1"]["subject"])

        # Body
        p = doc.add_paragraph()
        run = p.add_run('Treść:')
        run.bold = True

        for paragraph_text in camp["email_1"]["body"].split('\n\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.paragraph_format.space_after = Pt(6)

        wc = count_words(camp["email_1"]["body"])
        p = doc.add_paragraph()
        run = p.add_run(f'[{wc} słów]')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 7. Follow-up 1
        doc.add_heading('Follow-up 1 (step 2)', level=2)

        p = doc.add_paragraph()
        run = p.add_run('Temat: ')
        run.bold = True
        p.add_run(camp["follow_up_1"]["subject"])

        p = doc.add_paragraph()
        run = p.add_run('Treść:')
        run.bold = True

        for paragraph_text in camp["follow_up_1"]["body"].split('\n\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.paragraph_format.space_after = Pt(6)

        wc = count_words(camp["follow_up_1"]["body"])
        p = doc.add_paragraph()
        run = p.add_run(f'[{wc} słów]')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 8. Follow-up 2
        doc.add_heading('Follow-up 2 (step 3)', level=2)

        p = doc.add_paragraph()
        run = p.add_run('Temat: ')
        run.bold = True
        p.add_run(camp["follow_up_2"]["subject"])

        p = doc.add_paragraph()
        run = p.add_run('Treść:')
        run.bold = True

        for paragraph_text in camp["follow_up_2"]["body"].split('\n\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.paragraph_format.space_after = Pt(6)

        wc = count_words(camp["follow_up_2"]["body"])
        p = doc.add_paragraph()
        run = p.add_run(f'[{wc} słów]')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Page break between leads (except last)
        if idx < 2:
            doc.add_page_break()

    return doc


# ============================================================
# MAIN
# ============================================================

def main():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    # Ensure dirs exist
    os.makedirs(os.path.join(root, "tests", "output"), exist_ok=True)
    os.makedirs(os.path.join(root, "outputs", "word_campaigns"), exist_ok=True)

    # 1. Selected leads JSON
    path1 = os.path.join(root, "tests", "output", "tier2_real_pilot_selected_leads.json")
    with open(path1, "w", encoding="utf-8") as f:
        json.dump(generate_selected_leads(), f, ensure_ascii=False, indent=2)
    print(f"[OK] {path1}")

    # 2. Campaigns JSON
    path2 = os.path.join(root, "tests", "output", "tier2_real_pilot_campaigns.json")
    with open(path2, "w", encoding="utf-8") as f:
        json.dump(generate_campaigns(), f, ensure_ascii=False, indent=2)
    print(f"[OK] {path2}")

    # 3. Selection report MD
    path3 = os.path.join(root, "tests", "output", "tier2_real_pilot_selection_report.md")
    with open(path3, "w", encoding="utf-8") as f:
        f.write(generate_report())
    print(f"[OK] {path3}")

    # 4. Word document
    path4 = os.path.join(root, "outputs", "word_campaigns", "tier2_real_pilot_review.docx")
    doc = generate_word()
    doc.save(path4)
    print(f"[OK] {path4}")

    # Summary
    print("\n=== TIER 2 REAL PILOT - OUTPUTS GENERATED ===")
    print(f"Leads in CSV: {ALL_LEADS_COUNT}")
    print(f"Selected: 3")
    for lead in selected_leads:
        print(f"  {lead['id']}. {lead['first_name']} {lead['last_name']} ({lead['company']}) - {lead['angle']}")
    print(f"\nFiles:")
    print(f"  1. {path1}")
    print(f"  2. {path2}")
    print(f"  3. {path3}")
    print(f"  4. {path4}")
    print(f"\nStatus: READY FOR REVIEW (nie zsynchronizowano z Apollo)")


if __name__ == "__main__":
    main()
